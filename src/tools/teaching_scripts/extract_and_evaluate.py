#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
统一处理脚本：提取内容、查重检测、质量评估
适配新的 data/teaching 路径结构
"""

import os
import sys
import json

from tools.common import atomic_write_json
import re
from pathlib import Path
from docx import Document
from typing import List, Dict, Any, Optional

# 配置
BASE_DIR = Path(__file__).parent.parent.parent.parent
DATA_DIR = BASE_DIR / "data" / "teaching" / "2026-春季"


def extract_text_from_docx(docx_path):
    """Extract all text from a .docx file"""
    try:
        doc = Document(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n'.join(paragraphs), doc
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")
        return "", None


def extract_tables_from_docx(doc):
    """Extract tables from .docx document"""
    if not doc:
        return []
    try:
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables
    except Exception as e:
        print(f"Error reading tables: {e}")
        return []


def analyze_report_content(text, tables):
    """Analyze report content and extract key information"""
    content = {
        'word_count': len(text),
        'has_team_info': False,
        'has_objectives': False,
        'has_hardware_design': False,
        'has_software_design': False,
        'has_results': False,
        'has_discussion': False,
        'has_reflection': False,
        'team_members': [],
        'gpio_pins': [],
        'code_mentions': [],
        'key_findings': []
    }

    # Check for report sections
    keywords = {
        'has_team_info': ['团队成员', '小组成员', '分工', '成员'],
        'has_objectives': ['实验目的', '实验要求', '目标', '原理'],
        'has_hardware_design': ['硬件', '接线', '电路', 'GPIO', '引脚', '连接'],
        'has_software_design': ['软件', '程序', '代码', '流程', '状态机', '中断'],
        'has_results': ['测试', '结果', '现象', '演示', '效果'],
        'has_discussion': ['问题', '讨论', '分析', '解决'],
        'has_reflection': ['总结', '心得', '体会', '收获', '反思']
    }

    for key, words in keywords.items():
        for word in words:
            if word in text:
                content[key] = True
                break

    # Extract team members
    name_task_pattern = r'([^\s：:]+)\s*[：:]\s*(?:硬件|软件|接线|配置|中断|消抖|状态|报告|撰写|整体|DWT|LED|GPIO|EXTI|STM32)'
    matches = re.findall(name_task_pattern, text)
    if matches:
        exclude_words = ['说明', '请', '本', '本人', '小组', '团队', '记录', '成员']
        for match in matches:
            if match not in exclude_words and len(match) >= 2:
                content['team_members'].append(match)

    if '1.2 个人分工说明' in text or '分工说明' in text:
        content['has_team_info'] = True

    # Extract GPIO pin mentions
    gpio_pattern = r'(P[EF]\d+|PE\d+|PF\d+|GPIO|EXTI)'
    content['gpio_pins'] = list(set(re.findall(gpio_pattern, text)))

    # Look for code snippets or keywords
    code_keywords = ['HAL_GPIO', 'EXTI', 'NVIC', 'DWT', '中断', '回调', '消抖',
                    '状态', 'State', 'Gear', 'LED', 'KEY']
    for keyword in code_keywords:
        if keyword in text:
            content['code_mentions'].append(keyword)

    content['meets_word_requirement'] = content['word_count'] >= 1500

    return content


def process_student_report(student_data):
    """Process a single student's report"""
    student_id = student_data['id']
    student_name = student_data.get('name', '')
    docx_path = student_data.get('path', '')

    # 区分两种情况：
    # 1. 完全没有提交文件 (zip_name 为空)
    # 2. 提交了但没有 docx 文件（可能是图片或其他格式）
    has_submission = bool(student_data.get('zip_name'))
    has_docx = bool(docx_path)

    if not has_docx:
        if has_submission:
            print(f"  [跳过] {student_id}: 未按照模板提交（提交了非 DOCX 格式）")
            submission_status = 'wrong_format'  # 未按照模板提交
        else:
            print(f"  [跳过] {student_id}: 未提交报告")
            submission_status = 'missing'  # 完全未提交
        return {
            'student_id': student_id,
            'name': student_name,
            'docx_path': '',
            'full_text': '',
            'word_count': 0,
            'analysis': {'meets_word_requirement': False},
            'tables_count': 0,
            'submission_status': submission_status,
            'missing': not has_submission
        }

    print(f"  [处理] {student_id}: {student_name}")

    text, doc = extract_text_from_docx(docx_path)
    tables = extract_tables_from_docx(doc)

    analysis = analyze_report_content(text, tables)

    return {
        'student_id': student_id,
        'name': student_name,
        'docx_path': str(docx_path),
        'full_text': text[:20000],  # Increased for better plagiarism detection
        'word_count': analysis['word_count'],
        'analysis': analysis,
        'tables_count': len(tables),
        'missing': False
    }


def extract_content(class_name: str, experiment: str = "07-car-gear") -> List[Dict]:
    """
    Extract content from student reports

    Args:
        class_name: 班级名称
        experiment: 实验名称

    Returns:
        提取的内容列表
    """
    print("=" * 60)
    print(f"提取内容: {class_name} - {experiment}")
    print("=" * 60)

    experiment_dir = DATA_DIR / class_name / experiment
    processed_dir = experiment_dir / "processed"

    student_list_path = processed_dir / "students.json"
    if not student_list_path.exists():
        print(f"错误: 未找到 students.json: {student_list_path}")
        return []

    with open(student_list_path, 'r', encoding='utf-8') as f:
        students = json.load(f)

    results = []
    for student in students:
        result = process_student_report(student)
        results.append(result)

    output_path = processed_dir / "extracted_content.json"
    atomic_write_json(output_path, results, ensure_ascii=False, indent=2)

    print()
    print(f"完成! 提取了 {len(results)} 个学生的内容")
    print(f"结果保存至: {output_path}")

    avg_words = sum(r['word_count'] for r in results) / len(results) if results else 0
    print(f"平均字数: {avg_words:.0f}")
    print(f"达标学生: {sum(1 for r in results if r['analysis']['meets_word_requirement'])}/{len(results)}")
    print()

    return results


def run_plagiarism_detection(class_name: str, experiment: str = "07-car-gear",
                            threshold: float = 60.0) -> Dict:
    """
    Run plagiarism detection

    Args:
        class_name: 班级名称
        experiment: 实验名称
        threshold: 相似度阈值

    Returns:
        查重结果
    """
    print("=" * 60)
    print(f"查重检测: {class_name} - {experiment}")
    print(f"阈值: {threshold}%")
    print("=" * 60)

    experiment_dir = DATA_DIR / class_name / experiment
    processed_dir = experiment_dir / "processed"

    # Load extracted content
    content_path = processed_dir / "extracted_content.json"
    if not content_path.exists():
        print("错误: 未找到 extracted_content.json")
        return {}

    with open(content_path, 'r', encoding='utf-8') as f:
        extracted_data = json.load(f)

    # Import plagiarism detector
    try:
        sys.path.insert(0, str(BASE_DIR / "src" / "tools" / "plagiarism"))
        from core import PlagiarismDetector, SimilarityMethod
    except ImportError:
        try:
            from tools.plagiarism.core import PlagiarismDetector, SimilarityMethod
        except ImportError:
            print("错误: 无法导入 PlagiarismDetector")
            return {}

    # Build submissions dictionary
    submissions = {}
    for item in extracted_data:
        if item.get('missing'):
            continue
        submissions[item['student_id']] = {
            'name': item.get('name', ''),
            'text': item.get('full_text', '')
        }

    print(f"待检测报告数: {len(submissions)}")

    # Run detection
    detector = PlagiarismDetector(
        method=SimilarityMethod.HYBRID,
        threshold=threshold
    )

    all_results, suspicious, adaptive_report = detector.detect(submissions)

    # Save results
    output_path = processed_dir / "plagiarism_results.json"

    # Format results
    suspicious_data = []
    for r in suspicious:
        suspicious_data.append({
            'student_id': r.student_id,
            'student_name': submissions[r.student_id].get('name', ''),
            'similar_to': r.similar_to,
            'similar_name': submissions[r.similar_to].get('name', '') if r.similar_to in submissions else '',
            'overall_similarity': round(r.overall_similarity, 2),
            'is_cross_group': r.is_cross_group,
            'details': {
                'text_similarity': round(getattr(r, 'text_similarity', 0), 2),
                'code_similarity': round(getattr(r, 'code_similarity', 0), 2),
                'structure_similarity': round(getattr(r, 'structure_similarity', 0), 2),
                'semantic_similarity': round(getattr(r, 'semantic_similarity', 0), 2)
            }
        })

    results_data = {
        'total_students': len(submissions),
        'suspicious_count': len(suspicious),
        'threshold': threshold,
        'adaptive_report': adaptive_report,
        'suspicious_pairs': suspicious_data
    }

    atomic_write_json(output_path, results_data, ensure_ascii=False, indent=2)

    print()
    print(f"完成! 发现 {len(suspicious)} 组可疑报告")
    print(f"结果保存至: {output_path}")

    # Print top suspicious pairs
    if suspicious_data:
        print()
        print("相似度最高的前5组:")
        sorted_suspicious = sorted(suspicious_data, key=lambda x: -x['overall_similarity'])[:5]
        for item in sorted_suspicious:
            print(f"  {item['student_id']} ({item['student_name']}) <-> "
                  f"{item['similar_to']} ({item['similar_name']}): "
                  f"{item['overall_similarity']}%")
    print()

    return results_data


def run_quality_assessment(class_name: str, experiment: str = "07-car-gear") -> Dict:
    """
    Run quality assessment

    Args:
        class_name: 班级名称
        experiment: 实验名称

    Returns:
        质量评估结果
    """
    print("=" * 60)
    print(f"质量评估: {class_name} - {experiment}")
    print("=" * 60)

    experiment_dir = DATA_DIR / class_name / experiment
    processed_dir = experiment_dir / "processed"

    # Load extracted content
    content_path = processed_dir / "extracted_content.json"
    if not content_path.exists():
        print("错误: 未找到 extracted_content.json")
        return {}

    with open(content_path, 'r', encoding='utf-8') as f:
        extracted_data = json.load(f)

    # Load plagiarism results
    plagiarism_path = processed_dir / "plagiarism_results.json"
    plagiarism_data = {}
    if plagiarism_path.exists():
        with open(plagiarism_path, 'r', encoding='utf-8') as f:
            plagiarism_results = json.load(f)
            # Build plagiarism lookup
            for pair in plagiarism_results.get('suspicious_pairs', []):
                sid = pair['student_id']
                if sid not in plagiarism_data:
                    plagiarism_data[sid] = []
                plagiarism_data[sid].append(pair)

    # Score each student
    results = []
    for item in extracted_data:
        student_id = item['student_id']
        analysis = item.get('analysis', {})

        # Calculate scores
        scores = {
            'team_collaboration': 5 if analysis.get('has_team_info') else 2,
            'principle_understanding': min(10, (5 if analysis.get('has_objectives') else 2) +
                                          (5 if analysis.get('has_reflection') else 2)),
            'completion': min(35, (10 if analysis.get('has_hardware_design') else 0) +
                                 (15 if analysis.get('has_software_design') else 0) +
                                 (10 if analysis.get('has_results') else 0)),
            'code_quality': min(30, (15 if 'GPIO' in analysis.get('code_mentions', []) else 0) +
                                    (15 if 'EXTI' in analysis.get('code_mentions', []) else 0)),
            'report_quality': min(10, (5 if analysis.get('has_discussion') else 0) +
                                       (5 if analysis.get('meets_word_requirement') else 0))
        }

        total_score = sum(scores.values())

        # Check plagiarism
        plagiarism_info = plagiarism_data.get(student_id, [])
        is_plagiarism = any(p['overall_similarity'] >= 80 for p in plagiarism_info)

        # Apply plagiarism penalty
        final_score = 0 if is_plagiarism else total_score

        results.append({
            'student_id': student_id,
            'name': item.get('name', ''),
            'scores': scores,
            'total_score': total_score,
            'final_score': final_score,
            'is_plagiarism': is_plagiarism,
            'plagiarism_info': plagiarism_info,
            'word_count': item.get('word_count', 0),
            'submission_status': item.get('submission_status', 'submitted'),  # submitted, missing, wrong_format
            'missing': item.get('missing', False)
        })

    # Save results
    output_path = processed_dir / "quality_assessment.json"
    atomic_write_json(output_path, results, ensure_ascii=False, indent=2)

    print()
    print(f"完成! 评估了 {len(results)} 个学生")
    print(f"结果保存至: {output_path}")

    # Print statistics
    valid_scores = [r['final_score'] for r in results if not r['missing']]
    if valid_scores:
        avg_score = sum(valid_scores) / len(valid_scores)
        plagiarism_count = sum(1 for r in results if r['is_plagiarism'])
        print(f"平均分: {avg_score:.2f}")
        print(f"抄袭人数: {plagiarism_count}")
    print()

    return {'students': results}


def main():
    import argparse

    parser = argparse.ArgumentParser(description='统一处理脚本：提取内容、查重检测、质量评估')
    parser.add_argument('--class', dest='class_name', required=True, help='班级名称')
    parser.add_argument('--experiment', default='07-car-gear', help='实验名称')
    parser.add_argument('--skip-extract', action='store_true', help='跳过内容提取')
    parser.add_argument('--skip-plagiarism', action='store_true', help='跳过查重检测')
    parser.add_argument('--skip-quality', action='store_true', help='跳过质量评估')
    parser.add_argument('--threshold', type=float, default=60.0, help='查重阈值')

    args = parser.parse_args()

    print()
    print("=" * 60)
    print("教学评估系统 - 统一处理脚本")
    print(f"班级: {args.class_name}")
    print(f"实验: {args.experiment}")
    print("=" * 60)
    print()

    # Step 1: Extract content
    if not args.skip_extract:
        extract_content(args.class_name, args.experiment)

    # Step 2: Plagiarism detection
    if not args.skip_plagiarism:
        run_plagiarism_detection(args.class_name, args.experiment, args.threshold)

    # Step 3: Quality assessment
    if not args.skip_quality:
        run_quality_assessment(args.class_name, args.experiment)

    print()
    print("=" * 60)
    print("全部完成!")
    print("=" * 60)


if __name__ == "__main__":
    main()
