"""
从学生实验报告中提取内容
"""

import os
import json
import re
from pathlib import Path
from docx import Document

# Get the project root directory (go up from scripts to common to teaching to docs to project root)
SCRIPT_DIR = Path(__file__).parent
BASE_DIR = SCRIPT_DIR.parent.parent.parent.parent  # This should be the project root: NewProjectF407
EXPERIMENT_DIR = BASE_DIR / "docs" / "teaching" / "2026-春季" / "汽服2302B班" / "07-car-gear"
PROCESSED_DIR = EXPERIMENT_DIR / "processed"
RUBRIC_FILE = SCRIPT_DIR.parent / "rubrics" / "rubric.json"

def extract_text_from_docx(docx_path):
    """Extract all text from a .docx file"""
    try:
        doc = Document(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n'.join(paragraphs), doc
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")
        return "", None

def extract_tables_from_docx(doc):
    """Extract tables from .docx document"""
    if not doc:
        return []
    try:
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables
    except Exception as e:
        print(f"Error reading tables: {e}")
        return []

def analyze_report_content(text, tables):
    """Analyze report content and extract key information"""
    content = {
        'word_count': len(text),
        'has_team_info': False,
        'has_objectives': False,
        'has_hardware_design': False,
        'has_software_design': False,
        'has_results': False,
        'has_discussion': False,
        'has_reflection': False,
        'team_members': [],
        'gpio_pins': [],
        'code_mentions': [],
        'key_findings': []
    }

    # Check for report sections
    keywords = {
        'has_team_info': ['团队成员', '小组成员', '分工', '成员'],
        'has_objectives': ['实验目的', '实验要求', '目标', '原理'],
        'has_hardware_design': ['硬件', '接线', '电路', 'GPIO', '引脚', '连接'],
        'has_software_design': ['软件', '程序', '代码', '流程', '状态机', '中断'],
        'has_results': ['测试', '结果', '现象', '演示', '效果'],
        'has_discussion': ['问题', '讨论', '分析', '解决'],
        'has_reflection': ['总结', '心得', '体会', '收获', '反思']
    }

    for key, words in keywords.items():
        for word in words:
            if word in text:
                content[key] = True
                break

    # Extract team members (looking for patterns like 姓名：分工说明)
    # The reports have names followed by task descriptions
    # Pattern 1: 姓名：任务描述
    name_task_pattern = r'([^\s：:]+)\s*[：:]\s*(?:硬件|软件|接线|配置|中断|消抖|状态|报告|撰写|整体|DWT|LED|GPIO|EXTI|STM32)'
    matches = re.findall(name_task_pattern, text)
    if matches:
        # Filter out common non-name words
        exclude_words = ['说明', '请', '本', '本人', '小组', '团队', '记录', '成员']
        for match in matches:
            if match not in exclude_words and len(match) >= 2:
                content['team_members'].append(match)

    # Also check for explicit team section headers
    if '1.2 个人分工说明' in text or '分工说明' in text:
        content['has_team_info'] = True

    # Extract GPIO pin mentions
    gpio_pattern = r'(P[EF]\d+|PE\d+|PF\d+|GPIO|EXTI)'
    content['gpio_pins'] = list(set(re.findall(gpio_pattern, text)))

    # Look for code snippets or keywords
    code_keywords = ['HAL_GPIO', 'EXTI', 'NVIC', 'DWT', '中断', '回调', '消抖',
                    '状态', 'State', 'Gear', 'LED', 'KEY']
    for keyword in code_keywords:
        if keyword in text:
            content['code_mentions'].append(keyword)

    # Check for minimum word count
    content['meets_word_requirement'] = content['word_count'] >= 1500  # Approximate 2000 Chinese characters

    return content

def extract_name_from_text(text, doc):
    """Try to extract student name from report content"""
    if not text:
        return None

    # Pattern 1: Look for team member names in the text
    # Reports often have patterns like "张三负责...，李四负责..."
    name_patterns = [
        r'([一-龥]{2,3})负责',  # Name followed by 负责
        r'([一-龥]{2,3})[:：](?:硬件|软件|接线|配置|报告)',  # Name: task
        r'成员[：:]\s*([^\n]+)',  # Members: list
    ]

    for pattern in name_patterns:
        matches = re.findall(pattern, text)
        if matches:
            # Filter out common non-name words
            exclude = ['说明', '请', '本', '本人', '小组', '团队', '记录', '成员', '分工']
            for match in matches:
                if match and match not in exclude and len(match) >= 2:
                    # Extract just the first name found
                    clean_name = re.findall(r'([一-龥]{2,3})', match)
                    if clean_name:
                        return clean_name[0]

    # Pattern 2: Look for name in docx tables
    if doc:
        for table in doc.tables[:5]:  # Check first 5 tables
            for row in table.rows:
                row_text = ' '.join([cell.text for cell in row.cells])
                if '姓名' in row_text or '成员' in row_text:
                    # This row might contain names
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        # Look for Chinese names (2-3 characters)
                        name_match = re.search(r'([一-龥]{2,3})', cell_text)
                        if name_match and '姓名' not in cell_text and '成员' not in cell_text:
                            return name_match.group(1)

    return None

def process_student_report(student_data):
    """Process a single student's report"""
    student_id = student_data['id']
    student_name = student_data.get('name', '')
    docx_path = student_data.get('path', '')

    # Handle missing submissions
    if student_data.get('missing') or not docx_path:
        print(f"  Skipping {student_id}: No report submitted")
        return {
            'student_id': student_id,
            'name': student_name,
            'docx_path': '',
            'full_text': '',
            'word_count': 0,
            'analysis': {'meets_word_requirement': False},
            'tables_count': 0,
            'missing': True
        }

    print(f"  Processing {student_id}...")

    # Extract content
    text, doc = extract_text_from_docx(docx_path)
    tables = extract_tables_from_docx(doc)

    # Try to extract name from report content if not available from student list
    if not student_name:
        student_name = extract_name_from_text(text, doc) or ''

    # Analyze content
    analysis = analyze_report_content(text, tables)

    return {
        'student_id': student_id,
        'name': student_name,
        'docx_path': str(docx_path),
        'full_text': text[:15000],  # Increased to capture more sections
        'word_count': analysis['word_count'],
        'analysis': analysis,
        'tables_count': len(tables),
        'missing': False
    }

def main():
    print("Extracting content from student reports...")

    # Load student list
    student_list_path = PROCESSED_DIR / "students.json"
    if not student_list_path.exists():
        print("Error: students.json not found. Run main.py first.")
        return

    with open(student_list_path, 'r', encoding='utf-8') as f:
        students = json.load(f)

    # Process each student
    results = []
    for student in students:
        result = process_student_report(student)
        results.append(result)

    # Save results
    output_path = PROCESSED_DIR / "extracted_content.json"
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nExtracted content from {len(results)} students")
    print(f"Results saved to: {output_path}")

    # Print summary statistics
    avg_words = sum(r['word_count'] for r in results) / len(results)
    print(f"\nSummary:")
    print(f"  Average word count: {avg_words:.0f}")
    print(f"  Students meeting 2000+ words: {sum(1 for r in results if r['analysis']['meets_word_requirement'])}/{len(results)}")

if __name__ == "__main__":
    main()
