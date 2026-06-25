#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
反馈生成工作线程
Feedback Generation Worker Thread

把原先在 GUI 线程同步执行的「批量生成学生反馈 / 教师分析报告」（每个学生/班级
都要写 .md + .docx）移到后台线程，避免大批量时界面冻结。
"""

import os
import sys
import subprocess
from collections import OrderedDict
from pathlib import Path
from typing import Optional, List, Dict, Tuple

from PyQt6.QtCore import QThread, pyqtSignal

project_root = Path(__file__).parent.parent.parent.parent.parent
sys.path.insert(0, str(project_root))

from tools.teaching_management_gui.feedback_reports import (  # noqa: E402
    build_student_feedback,
    build_teacher_report,
    build_group_feedback,
    pick_group_leader,
)
from tools.teaching_management_gui.path_helper import (  # noqa: E402
    feedback_dir as resolve_feedback_dir,
)

try:  # noqa: E402
    from docx import Document as _DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _write_docx(path: Path, title: str, text: str):
    """把多行文本写成 Word（标题 + 列表/段落）。"""
    doc = _DocxDocument()
    doc.add_heading(title, level=1)
    for line in text.splitlines():
        if line.startswith("# "):
            continue  # 主标题已加
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("| "):
            doc.add_paragraph(line)  # 表格行按段落保留
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(path)


def generate_group_feedback_files(
    reports: List[Tuple[Dict, str, str]],
    semester: str,
    include_strengths: bool = True,
    include_weaknesses: bool = True,
    include_suggestions: bool = True,
    concise: bool = False,
    log=print,
    on_progress=None,
    is_cancelled=None,
) -> Tuple[int, int]:
    """按 ``(班级, 实验, group_key)`` 聚合，每组生成 1 份反馈。

    - 同组（共享同一份小组报告/工程）的成员合并为一份，列出全体组员与各自成绩；
    - 无 ``group_key``（个人实验或旧数据）按学号自成一组，行为退化为每生一份；
    - md → ``学生反馈/md/``，word → ``学生反馈/word/``（分文件夹）。

    Args:
        reports: [(report_dict, class_name, experiment_id), ...]
        semester: 学期（决定产物路径）
        on_progress: 可选回调 (current, total) -> None
        is_cancelled: 可选取消判定 () -> bool

    Returns:
        (success, total)
    """
    groups: "OrderedDict[Tuple[str, str, str], List[Dict]]" = OrderedDict()
    for rep, cls, exp in reports:
        gk = rep.get('group_key') or rep.get('student_id') or 'unknown'
        groups.setdefault((cls, exp, gk), []).append(rep)

    total = len(groups)
    success = 0
    log(f"批量生成小组反馈：{total} 个小组（来源 {len(reports)} 份个人报告）")

    for i, ((cls, exp, gk), member_reports) in enumerate(groups.items()):
        if on_progress:
            on_progress(i, total)
        if is_cancelled and is_cancelled():
            break
        try:
            text = build_group_feedback(
                member_reports, cls, exp,
                include_strengths, include_weaknesses, include_suggestions, concise,
            )
            base_dir = resolve_feedback_dir(cls, exp, semester) / "学生反馈"
            md_dir = base_dir / "md"
            word_dir = base_dir / "word"
            md_dir.mkdir(parents=True, exist_ok=True)
            word_dir.mkdir(parents=True, exist_ok=True)

            leader = pick_group_leader(member_reports)
            lsid = leader.get('student_id', '') or gk
            lname = leader.get('name', '')
            n = len(member_reports)
            base = f"{lsid}_{lname}_小组反馈({n}人)" if lname else f"{gk}_小组反馈({n}人)"

            (md_dir / f"{base}.md").write_text(text, encoding="utf-8")
            if HAS_DOCX:
                _write_docx(word_dir / f"{base}.docx",
                            f"{lname or gk} 小组实验反馈（{n}人）", text)
            success += 1
        except Exception as e:
            log(f"生成失败 小组 {gk}: {e}")

    if on_progress:
        on_progress(total, total)
    log(f"小组反馈完成：{success}/{total}")
    return success, total


class FeedbackWorker(QThread):
    """批量反馈/报告生成线程。

    输入与 FeedbackPanel.generate_feedback 收集的一致；在 run() 内完成全部文件写入，
    通过信号回报进度/日志/结果，GUI 线程不参与磁盘与 docx 渲染。
    """

    log_message = pyqtSignal(str)
    progress = pyqtSignal(int, int)            # (current, total)
    feedback_completed = pyqtSignal(int, int)  # (success, total)
    feedback_failed = pyqtSignal(str)
    feedback_cancelled = pyqtSignal()

    def __init__(
        self,
        mode: str,
        reports,
        semester: str,
        include_strengths: bool = True,
        include_weaknesses: bool = True,
        include_suggestions: bool = True,
        concise: bool = False,
    ):
        super().__init__()
        self.mode = mode  # "student" | "teacher"
        self.reports = list(reports)
        self.semester = semester
        self.include_strengths = include_strengths
        self.include_weaknesses = include_weaknesses
        self.include_suggestions = include_suggestions
        self.concise = concise
        self.is_cancelled = False

    def cancel(self):
        self.is_cancelled = True
        self.log_message.emit("正在取消...")

    def run(self):
        try:
            if not HAS_DOCX:
                self.log_message.emit("提示：未安装 python-docx，仅生成 Markdown。")

            if self.mode == "teacher":
                success, total = self._generate_teacher()
            else:
                success, total = self._generate_student()

            if self.is_cancelled:
                self.log_message.emit(f"反馈生成已取消（已完成 {success}/{total}）")
                self.feedback_cancelled.emit()
            else:
                # 生成成功后自动在文件资源管理器中打开输出目录
                if success > 0:
                    opened = self._open_output_folder()
                    if opened:
                        self.log_message.emit(f"已打开输出目录：{opened}")
                self.feedback_completed.emit(success, total)
        except Exception as e:
            self.feedback_failed.emit(str(e))

    def _open_output_folder(self) -> Optional[Path]:
        """生成完成后打开输出目录（Windows 资源管理器）。返回打开的路径，失败返回 None。

        多班级时打开第一个班级的目录；学生模式开「学生反馈」，教师模式开「教师报告」。
        """
        if not self.reports:
            return None
        _, cls, exp = self.reports[0]
        sub = "教师报告" if self.mode == "teacher" else "学生反馈"
        folder = resolve_feedback_dir(cls, exp, self.semester) / sub
        if not folder.exists():
            folder = folder.parent
        if not folder.exists():
            return None
        try:
            if sys.platform.startswith("win"):
                os.startfile(str(folder))  # type: ignore[attr-defined]
            elif sys.platform == "darwin":
                subprocess.Popen(["open", str(folder)])
            else:
                subprocess.Popen(["xdg-open", str(folder)])
        except Exception as e:
            self.log_message.emit(f"（无法自动打开文件夹：{e}）")
            return None
        return folder

    # ---------------- 学生反馈（按小组聚合，每组 1 份） ----------------
    def _generate_student(self):
        success, total = generate_group_feedback_files(
            self.reports,
            self.semester,
            self.include_strengths,
            self.include_weaknesses,
            self.include_suggestions,
            self.concise,
            log=self.log_message.emit,
            on_progress=lambda c, t: self.progress.emit(c, t),
            is_cancelled=lambda: self.is_cancelled,
        )
        if self.is_cancelled:
            self.log_message.emit(f"反馈生成已取消（已完成 {success}/{total}）")
        return success, total

    # ---------------- 教师分析报告 ----------------
    def _generate_teacher(self):
        # 按 (class, exp) 去重保序
        classes = []
        for _r, cls, exp in self.reports:
            if (cls, exp) not in classes:
                classes.append((cls, exp))
        total = len(classes)
        success = 0
        self.log_message.emit(f"批量生成教师分析报告：{total} 个班级")
        for i, (cls, exp) in enumerate(classes):
            if self.is_cancelled:
                break
            self.progress.emit(i, total)
            cls_reports = [r for r, c, _e in self.reports if c == cls]
            if not cls_reports:
                continue
            try:
                text = build_teacher_report(cls, exp, cls_reports)
                out_dir = resolve_feedback_dir(cls, exp, self.semester) / "教师报告"
                out_dir.mkdir(parents=True, exist_ok=True)
                base = f"{cls}_{exp}_教师分析报告"
                (out_dir / f"{base}.md").write_text(text, encoding="utf-8")
                if HAS_DOCX:
                    _write_docx(out_dir / f"{base}.docx", f"{cls} 教学分析报告", text)
                success += 1
                self.log_message.emit(f"  {cls}：已生成（{len(cls_reports)} 人）")
            except Exception as e:
                self.log_message.emit(f"  {cls} 生成失败: {e}")
        self.progress.emit(total, total)
        if not self.is_cancelled:
            self.log_message.emit(f"教师分析报告完成：{success}/{total}")
        return success, total
