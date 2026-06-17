#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
反馈生成面板（批量版）
Feedback Generation Panel

输入取自「数据源」页（多班级），遍历所有已选班级读取批阅结果并生成反馈。
"""

import sys
import json
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent.parent))

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout,
    QGroupBox, QLabel, QPushButton,
    QComboBox, QCheckBox, QTextEdit, QMessageBox,
)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QFont

from tools.teaching_management_gui.data_source import shared
from tools.teaching_management_gui.path_helper import (
    grading_dir as resolve_grading_dir,
    feedback_dir as resolve_feedback_dir,
)

try:
    from docx import Document as _DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class FeedbackPanel(QWidget):
    """反馈生成面板"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.is_generating = False
        self.reports = []   # [(report_dict, class_name), ...]

        self.setup_ui()
        shared().entries_changed.connect(self._refresh_data_source_status)
        self._refresh_data_source_status()

    # ---------------- UI ----------------
    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.addWidget(self._create_config_panel())
        layout.addWidget(self._create_preview_panel(), 1)
        layout.addWidget(self._create_log_panel())

    def _create_config_panel(self):
        group = QGroupBox("数据源与反馈参数")
        v = QVBoxLayout()

        ds_row = QHBoxLayout()
        self.ds_status = QLabel()
        self.ds_status.setWordWrap(True)
        ds_row.addWidget(self.ds_status, 1)
        go_btn = QPushButton("前往数据源页")
        go_btn.clicked.connect(lambda: self.window().nav_list.setCurrentRow(0))
        ds_row.addWidget(go_btn)
        v.addLayout(ds_row)

        row = QHBoxLayout()
        row.addWidget(QLabel("反馈类型:"))
        self.feedback_type = QComboBox()
        self.feedback_type.addItem("详细反馈（含建议）")
        self.feedback_type.addItem("简洁反馈")
        self.feedback_type.addItem("成绩单")
        row.addWidget(self.feedback_type)
        row.addStretch()
        v.addLayout(row)

        opts = QHBoxLayout()
        self.include_strengths = QCheckBox("包含优点")
        self.include_strengths.setChecked(True)
        self.include_weaknesses = QCheckBox("包含不足")
        self.include_weaknesses.setChecked(True)
        self.include_suggestions = QCheckBox("包含改进建议")
        self.include_suggestions.setChecked(True)
        opts.addWidget(self.include_strengths)
        opts.addWidget(self.include_weaknesses)
        opts.addWidget(self.include_suggestions)
        opts.addStretch()
        v.addLayout(opts)

        btn_row = QHBoxLayout()
        btn_row.addStretch()
        self.preview_btn = QPushButton("预览反馈")
        self.preview_btn.setMinimumHeight(38)
        self.preview_btn.clicked.connect(self.preview_feedback)
        btn_row.addWidget(self.preview_btn)
        self.generate_btn = QPushButton("批量生成（全部班级）")
        self.generate_btn.setMinimumHeight(38)
        self.generate_btn.setStyleSheet(
            "QPushButton{background-color:#FF9800;color:white;font-weight:bold;"
            "border-radius:5px;padding:8px 16px;}"
            "QPushButton:hover{background-color:#F57C00;}"
            "QPushButton:disabled{background-color:#cccccc;}"
        )
        self.generate_btn.clicked.connect(self.generate_feedback)
        btn_row.addWidget(self.generate_btn)
        btn_row.addStretch()
        v.addLayout(btn_row)

        group.setLayout(v)
        return group

    def _create_preview_panel(self):
        group = QGroupBox("反馈预览")
        v = QVBoxLayout()
        row = QHBoxLayout()
        row.addWidget(QLabel("选择学生:"))
        self.student_combo = QComboBox()
        self.student_combo.addItem("请先点击「预览反馈」加载…")
        row.addWidget(self.student_combo, 1)
        v.addLayout(row)
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        self.preview_text.setFont(QFont("Microsoft YaHei", 10))
        self.preview_text.setPlaceholderText("点击「预览反馈」查看内容…")
        v.addWidget(self.preview_text)
        group.setLayout(v)
        return group

    def _create_log_panel(self):
        group = QGroupBox("日志输出")
        group.setMaximumHeight(120)
        v = QVBoxLayout()
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setFont(QFont("Consolas", 9))
        v.addWidget(self.log_text)
        group.setLayout(v)
        return group

    # ---------------- 数据源状态 ----------------
    def _refresh_data_source_status(self):
        entries = shared().entries()
        if not entries:
            self.ds_status.setText("⚠ 尚未选择班级压缩包。请到「数据源」页选择（可多选）。")
            self.ds_status.setStyleSheet("color:#c0392b;font-weight:bold;")
            self.generate_btn.setEnabled(False)
            self.preview_btn.setEnabled(False)
        else:
            names = "、".join(e.class_name for e in entries)
            self.ds_status.setText(f"✓ 已选 {len(entries)} 个班级：{names}")
            self.ds_status.setStyleSheet("color:#27ae60;font-weight:bold;")
            self.generate_btn.setEnabled(not self.is_generating)
            self.preview_btn.setEnabled(not self.is_generating)

    def log(self, message):
        if hasattr(self, "log_text"):
            ts = datetime.now().strftime("%H:%M:%S")
            self.log_text.append(f"[{ts}] {message}")

    # ---------------- 加载 / 构建 ----------------
    def _load_all_reports(self):
        """跨班级加载所有个人报告，返回 [(report_dict, class_name), ...]。"""
        semester = shared().semester()
        reports = []
        for e in shared().entries():
            individuals = resolve_grading_dir(e.class_name, e.experiment_id, semester) / "个人报告"
            if not individuals.exists():
                self.log(f"未找到个人报告：{individuals}（{e.class_name}，请先在「评分」中批阅）")
                continue
            for f in sorted(individuals.glob("*-评分.json")):
                try:
                    reports.append((json.loads(f.read_text(encoding="utf-8")), e.class_name))
                except Exception as ex:
                    self.log(f"读取失败 {f.name}: {ex}")
        self.reports = reports
        # 刷新下拉（带班级前缀）
        self.student_combo.blockSignals(True)
        self.student_combo.clear()
        for rep, cls in reports:
            self.student_combo.addItem(
                f"[{cls}] {rep.get('name', '?')}({rep.get('student_id', '?')}) - {rep.get('grade', '')}"
            )
        self.student_combo.blockSignals(False)
        return reports

    def _build_feedback_text(self, report: dict, class_name: str = "") -> str:
        ftype = self.feedback_type.currentText()
        inc_s = self.include_strengths.isChecked()
        inc_w = self.include_weaknesses.isChecked()
        inc_g = self.include_suggestions.isChecked()

        name = report.get("name", "同学")
        sid = report.get("student_id", "")
        total = report.get("total_score", 0)
        max_score = report.get("max_score", 100)
        grade = report.get("grade", "N/A")

        lines = [
            f"{name} 同学（学号 {sid}，{class_name}）：",
            "",
            f"本次实验批阅结果：总分 {total}/{max_score}（等级 {grade}）",
        ]
        cat_scores = report.get("category_scores", [])
        if ftype == "成绩单":
            lines.extend(["", "【各项得分】"])
            for cs in cat_scores:
                lines.append(f"- {cs.get('category_name', cs.get('category_id', ''))}："
                             f"{cs.get('earned_points', 0)}/{cs.get('max_points', 0)}")
            return "\n".join(lines)

        if ftype.startswith("详细"):
            lines.extend(["", "【各项得分】"])
            for cs in cat_scores:
                lines.append(f"- {cs.get('category_name', cs.get('category_id', ''))}："
                             f"{cs.get('earned_points', 0)}/{cs.get('max_points', 0)}")

        strengths = report.get("strengths", []) or []
        weaknesses = report.get("weaknesses", []) or []
        suggestions = report.get("suggestions", []) or []
        if ftype.startswith("简洁"):
            strengths, weaknesses, suggestions = strengths[:2], weaknesses[:2], suggestions[:2]
        if inc_s and strengths:
            lines.extend(["", "【优点】", *[f"- {s}" for s in strengths]])
        if inc_w and weaknesses:
            lines.extend(["", "【不足】", *[f"- {s}" for s in weaknesses]])
        if inc_g and suggestions:
            lines.extend(["", "【改进建议】", *[f"- {s}" for s in suggestions]])
        lines.extend(["", "祝你学习进步！"])
        return "\n".join(lines)

    # ---------------- 预览 / 生成 ----------------
    def preview_feedback(self):
        try:
            reports = self._load_all_reports()
        except Exception as e:
            QMessageBox.warning(self, "提示", str(e))
            return
        if not reports:
            QMessageBox.warning(self, "提示", "未读取到任何个人报告，请先在「评分」中批阅。")
            return
        idx = max(self.student_combo.currentIndex(), 0)
        rep, cls = reports[idx]
        self.preview_text.setPlainText(self._build_feedback_text(rep, cls))
        self.log(f"已预览 [{cls}] {rep.get('name', '')} 的反馈")

    def generate_feedback(self):
        semester = shared().semester()
        entries = shared().entries()
        if not entries:
            QMessageBox.warning(self, "警告", "请先到「数据源」页选择班级")
            return
        try:
            self._load_all_reports()
        except Exception as e:
            QMessageBox.warning(self, "提示", str(e))
            return
        if not self.reports:
            QMessageBox.warning(self, "提示", "未读取到任何个人报告，请先在「评分」中批阅。")
            return

        self.is_generating = True
        self.generate_btn.setEnabled(False)
        self.preview_btn.setEnabled(False)
        self.log(f"开始批量生成反馈：覆盖 {len(entries)} 个班级、{len(self.reports)} 名学生")

        success = 0
        for rep, cls in self.reports:
            try:
                text = self._build_feedback_text(rep, cls)
                sid = rep.get("student_id", "unknown")
                name = rep.get("name", "")
                # 该生所在实验/班级对应的反馈目录
                exp_id = next((e.experiment_id for e in entries if e.class_name == cls), "")
                out_dir = resolve_feedback_dir(cls, exp_id, semester)
                out_dir.mkdir(parents=True, exist_ok=True)
                base = f"{sid}_{name}_反馈" if name else f"{sid}_反馈"
                (out_dir / f"{base}.md").write_text(text, encoding="utf-8")
                if HAS_DOCX:
                    doc = _DocxDocument()
                    doc.add_heading(f"{name or sid} 实验反馈", level=1)
                    for line in text.splitlines():
                        if line.startswith("- "):
                            doc.add_paragraph(line[2:], style="List Bullet")
                        elif line.startswith("【") and line.endswith("】"):
                            doc.add_heading(line.strip("【】"), level=2)
                        elif line.strip():
                            doc.add_paragraph(line)
                    doc.save(out_dir / f"{base}.docx")
                success += 1
            except Exception as e:
                self.log(f"生成失败 {rep.get('name', '')}: {e}")

        self.log(f"反馈生成完成：{success}/{len(self.reports)}")
        if not HAS_DOCX:
            self.log("提示：未安装 python-docx，仅生成 Markdown。")
        self.is_generating = False
        self.generate_btn.setEnabled(True)
        self.preview_btn.setEnabled(True)
        QMessageBox.information(self, "完成", f"已生成 {success}/{len(self.reports)} 份反馈")

    def export_report(self):
        """供主窗口「导出报告」菜单调用。"""
        self.generate_feedback()
