"""
将 Markdown 文件转换为 DOCX 格式
改进列表识别，避免误判
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from pathlib import Path


def set_chinese_font(run, font_name='宋体', size=None, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    if bold:
        run.font.bold = True


def parse_inline_formatting(text):
    """解析行内格式（加粗、斜体）"""
    parts = []
    remaining = text

    while remaining:
        # 加粗 **text**
        bold_match = re.match(r'(.*?)\*\*([^*]+?)\*\*(.*)', remaining)
        if bold_match:
            if bold_match.group(1):
                parts.append({'text': bold_match.group(1), 'bold': False})
            parts.append({'text': bold_match.group(2), 'bold': True})
            remaining = bold_match.group(3)
            continue

        # 斜体 *text*
        italic_match = re.match(r'(.*?)\*([^*]+?)\*(.*)', remaining)
        if italic_match and not italic_match.group(2).startswith('*'):
            if italic_match.group(1):
                parts.append({'text': italic_match.group(1), 'italic': False})
            parts.append({'text': italic_match.group(2), 'italic': True})
            remaining = italic_match.group(3)
            continue

        # 剩余文本
        parts.append({'text': remaining})
        break

    return parts


def add_formatted_paragraph(doc, text, style='Normal', line_spacing=1.35, first_line_indent=False):
    """添加带格式的段落"""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # 首行缩进2字符（约0.75cm）
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.75)

    # 解析格式
    parts = parse_inline_formatting(text)

    for part in parts:
        run = p.add_run(part.get('text', ''))
        set_chinese_font(run, '宋体', 12)
        if part.get('bold'):
            run.font.bold = True
        if part.get('italic'):
            run.font.italic = True

    return p


def is_list_item(text, in_list_context=False):
    """判断是否为列表项（更严格的判断）"""
    stripped = text.strip()

    # 只有以明确的列表标记开头才是列表
    # 排除：纯数字开头、章节标题、说明行等
    list_patterns = [
        r'^[\s]*[\-\*]\s+\S',      # - 或 * 开头
        r'^[\s]*\d+\.\s+\S',       # 数字. 开头（但后面必须有空格和内容）
        r'^[\s]*[\d]+、\s+\S', # 数字、开头（中文句号）
    ]

    for pattern in list_patterns:
        if re.match(pattern, stripped):
            # 排除一些特殊情况
            if re.match(r'^[\d]+\.\s*(说明|注|图|表)', stripped):
                return False
            if re.match(r'^[一二三四五六七八九十]、', stripped):
                return False
            if re.match(r'^\d+\.\s+\d+\.\s+', stripped):  # 排除 "1. 2." 这样的格式
                return False
            return True

    return False


def process_table_cell_text(text):
    """处理表格单元格中的格式"""
    parts = parse_inline_formatting(text)
    return parts


def md_to_docx(md_path, docx_path):
    """将 Markdown 文件转换为 DOCX"""
    md_path = Path(md_path)
    docx_path = Path(docx_path)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    i = 0
    in_code_block = False
    code_lines = []
    in_list = False
    list_type = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过适用学期和编制日期行
        if '适用学期' in stripped or ('编制日期' in stripped and '版本' not in stripped):
            i += 1
            continue

        # 空行处理
        if not stripped:
            if not in_code_block:
                in_list = False
            else:
                code_lines.append('')
            i += 1
            continue

        # 代码块处理
        if stripped.startswith('```'):
            if in_code_block:
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    shading_el = OxmlElement('w:shd')
                    shading_el.set(qn('w:fill'), 'F5F5F5')
                    p._element.get_or_add_pPr().append(shading_el)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # 分割线
        if stripped == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('—' * 40)
            set_chinese_font(run, '黑体', 10)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            in_list = False
            i += 1
            continue

        # 标题处理（支持 # 和 一二三四）
        if stripped.startswith('#'):
            match = re.match(r'^(#+)\s+(.+)', stripped)
            if match:
                level = min(len(match.group(1)), 3)
                text = match.group(2).strip()

                heading = doc.add_heading('', level)

                parts = parse_inline_formatting(text)
                for part in parts:
                    run = heading.add_run(part.get('text', ''))
                    size = {0: 18, 1: 16, 2: 14, 3: 13}.get(level, 12)
                    set_chinese_font(run, '黑体', size)
                    if part.get('bold'):
                        run.font.bold = True

                heading.paragraph_format.line_spacing = 1.5
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)

                # 大标题（Heading 1，level=1）居中
                if level == 1:
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                in_list = False
                i += 1
                continue

        # 表格处理
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            rows = []
            for tl in table_lines:
                if '|---' not in tl:
                    cells = [c.strip() for c in tl.split('|')]
                    if len(cells) >= 2:
                        cells = cells[1:-1] if tl.startswith('|') else cells
                        rows.append(cells)

            if rows and all(len(r) == len(rows[0]) for r in rows):
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'

                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        cell = table.rows[ri].cells[ci]
                        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

                        for para in cell.paragraphs:
                            for run in para.runs:
                                run._element.get_parent().get_parent().remove(run._element)

                        parts = process_table_cell_text(cell_text)
                        if not parts:
                            cell.text = cell_text
                            for para in cell.paragraphs:
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                para.paragraph_format.line_spacing = 1.25
                                for run in para.runs:
                                    if ri == 0:
                                        set_chinese_font(run, '黑体', 11, bold=True)
                                    else:
                                        set_chinese_font(run, '宋体', 11)
                        else:
                            para = cell.paragraphs[0]
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            para.paragraph_format.line_spacing = 1.25

                            for part in parts:
                                run = para.add_run(part.get('text', ''))
                                if ri == 0:
                                    set_chinese_font(run, '黑体', 11, bold=True)
                                    if part.get('bold'):
                                        run.font.bold = True
                                else:
                                    set_chinese_font(run, '宋体', 11)
                                    if part.get('bold'):
                                        run.font.bold = True

            in_list = False
            continue

        # 列表处理（严格判断）
        if is_list_item(stripped):
            # 使用普通段落样式，手动添加列表标记
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.line_spacing = 1.35
            p.paragraph_format.first_line_indent = Cm(0.75)

            # 提取列表标记和内容
            match = re.match(r'^\s*([\d\-\*]+[\.、]?)\s*(.*)', stripped)
            if match:
                list_mark = match.group(1)
                text = match.group(2)

                # 添加列表标记（带加粗）
                run = p.add_run(list_mark + ' ')
                set_chinese_font(run, '宋体', 12, bold=True)

                # 添加列表内容
                parts = parse_inline_formatting(text)
                for part in parts:
                    run = p.add_run(part.get('text', ''))
                    set_chinese_font(run, '宋体', 12)
                    if part.get('bold'):
                        run.font.bold = True
            else:
                # 如果没有匹配到，保留原文本
                parts = parse_inline_formatting(stripped)
                for part in parts:
                    run = p.add_run(part.get('text', ''))
                    set_chinese_font(run, '宋体', 12)
                    if part.get('bold'):
                        run.font.bold = True

            in_list = True
            i += 1
            continue

        # 普通段落
        # 检查是否为章节标题（中文数字）
        is_chapter = re.match(r'^[一二三四五六七八九十]+、', stripped)
        # 检查是否为说明行
        is_note = stripped.startswith('说明') or stripped.startswith('>')

        # 章节标题和说明不需要首行缩进
        need_indent = not (is_chapter or is_note)

        p = add_formatted_paragraph(doc, stripped, line_spacing=1.35, first_line_indent=need_indent)

        if is_chapter:
            p.paragraph_format.space_before = Pt(6)
            p.runs[0].font.bold = True
            set_chinese_font(p.runs[0], '黑体', 13)
        elif is_note:
            p.paragraph_format.space_before = Pt(3)

        in_list = False
        i += 1

    doc.save(docx_path)
    print(f'OK: {md_path.name} -> {docx_path.name}')
    return docx_path


def main():
    """转换指定的文件"""
    base_dir = Path(__file__).parent.parent

    print('正在转换文件（改进列表识别）...')
    md_to_docx(
        base_dir / 'docs' / '07_car_gear_experiment.md',
        base_dir / 'docs' / '07_car_gear_experiment.docx'
    )

    print('\n转换完成！')
    print('改进：')
    print('  - 严格列表识别，减少误判')
    print('  - 章节标题不作为列表')
    print('  - 说明文字不缩进')
    print('  - 保持加粗和格式')


if __name__ == '__main__':
    main()
