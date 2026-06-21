#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
多班级报告生成器
Multi-Class Report Generator

生成多班级查重检测的汇总报告，支持：
1. Excel多工作表报告
2. JSON结构化报告
3. HTML可视化报告

作者: STM32F407 教学团队
版本: 1.0.0
"""

import json
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
from dataclasses import dataclass

logger = logging.getLogger(__name__)

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.drawing.image import Image
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import matplotlib
    matplotlib.use('Agg')  # 非GUI后端
    import matplotlib.pyplot as plt
    import matplotlib
    from matplotlib import rcParams
    # 设置中文字体
    rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS']
    rcParams['axes.unicode_minus'] = False
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


@dataclass
class ReportConfig:
    """报告配置"""
    output_dir: Path
    project_name: str = "多班级查重"
    enable_color_coding: bool = True


class MultiClassReportGenerator:
    """多班级报告生成器"""

    def __init__(
        self,
        output_dir: Path,
        project_name: str = "多班级查重"
    ):
        """
        初始化报告生成器

        Args:
            output_dir: 输出目录
            project_name: 项目名称
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.project_name = project_name

        # 颜色配置
        self.color_high = "FFFF0000"      # 红色 - 高风险
        self.color_medium = "FFFF9900"    # 橙色 - 中风险
        self.color_low = "FF00FF00"       # 绿色 - 低风险
        self.color_header = "FF4472C4"    # 蓝色 - 表头

        # 评分数据
        self.grading_data: Dict[str, List[Dict]] = {}  # class_id -> grading results

    def generate_all(
        self,
        detection_result,
        formats: Optional[List[str]] = None
    ) -> List[Path]:
        """
        生成所有格式的报告

        Args:
            detection_result: MultiClassDetectionResult
            formats: 报告格式列表 ['excel', 'json', 'html']

        Returns:
            生成的报告路径列表
        """
        if formats is None:
            formats = ['excel', 'json']

        paths = []

        if 'excel' in formats and HAS_OPENPYXL:
            try:
                path = self.generate_excel(detection_result)
                paths.append(path)
            except Exception:
                logger.exception("Excel 报告生成失败")

        if 'json' in formats:
            try:
                path = self.generate_json(detection_result)
                paths.append(path)
            except Exception:
                logger.exception("JSON 报告生成失败")

        if 'pdf' in formats and HAS_REPORTLAB:
            try:
                path = self.generate_pdf(detection_result)
                paths.append(path)
            except Exception:
                logger.exception("PDF 报告生成失败")

        if 'word' in formats and HAS_PYTHON_DOCX:
            try:
                path = self.generate_word(detection_result)
                paths.append(path)
            except Exception:
                logger.exception("Word 报告生成失败")

        return paths

    def generate_excel(self, detection_result) -> Path:
        """生成Excel多工作表报告"""
        wb = openpyxl.Workbook()
        wb.remove(wb.active)  # 删除默认工作表

        # 1. 汇总表
        self._create_summary_sheet(wb, detection_result)

        # 2. 班级概览表
        self._create_class_overview_sheet(wb, detection_result)

        # 3. 跨班级可疑对表
        self._create_cross_class_sheet(wb, detection_result)

        # 4. 班级对比表
        self._create_comparison_sheet(wb, detection_result)

        # 5. 评分对比表（如果有评分数据）
        if self.grading_data:
            self._create_grading_comparison_sheet(wb, detection_result)

        # 6. 可视化图表（如果有matplotlib）
        if HAS_MATPLOTLIB and self.grading_data:
            chart_path = self._generate_comparison_charts(detection_result)
            # 在第一个工作表中插入图表

        # 7. 各班级详细表（每个班级一个工作表）
        for class_id, class_result in detection_result.class_results.items():
            self._create_class_detail_sheet(wb, class_id, class_result)

        # 保存
        filename = f"{self.project_name}_多班级汇总_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        output_path = self.output_dir / filename
        wb.save(output_path)

        return output_path

    def _create_summary_sheet(self, wb, detection_result):
        """创建汇总工作表"""
        ws = wb.create_sheet("汇总")

        # 标题
        ws['A1'] = f"{self.project_name} - 多班级查重汇总"
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:C1')

        # 生成时间
        ws['A2'] = f"生成时间: {detection_result.timestamp}"
        ws['A2'].font = Font(size=10)

        # 汇总统计
        summary = detection_result.get_summary()

        row = 4
        headers = ["统计项", "数值"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row, col, header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color=self.color_header, end_color=self.color_header, fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        row += 1
        stats = [
            ("班级数量", summary['total_classes']),
            ("总学生数", summary['total_students']),
            ("班级内可疑对", summary['total_suspicious_pairs']),
            ("跨班级可疑对", summary['cross_class_suspicious_pairs']),
            ("班级对比数", summary['class_comparisons']),
        ]

        for label, value in stats:
            ws.cell(row, 1, label)
            ws.cell(row, 2, value)
            row += 1

        # 调整列宽
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 15

    def _create_class_overview_sheet(self, wb, detection_result):
        """创建班级概览工作表"""
        ws = wb.create_sheet("班级概览")

        row = 1
        headers = ["班级ID", "班级名称", "学生数", "可疑对数", "可疑率", "团伙数"]

        # 表头
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row, col, header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color=self.color_header, end_color=self.color_header, fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        # 数据行
        for class_id, result in detection_result.class_results.items():
            row += 1
            ws.cell(row, 1, class_id)
            ws.cell(row, 2, result.class_name)
            ws.cell(row, 3, result.student_count)
            ws.cell(row, 4, result.suspicious_pairs)

            suspicious_rate = result.suspicious_pairs / result.student_count if result.student_count > 0 else 0
            ws.cell(row, 5, f"{suspicious_rate:.1%}")
            ws.cell(row, 6, len(result.groups))

        # 调整列宽
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 10
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 10
        ws.column_dimensions['F'].width = 10

    def _create_cross_class_sheet(self, wb, detection_result):
        """创建跨班级可疑对工作表"""
        ws = wb.create_sheet("跨班级可疑对")

        row = 1
        headers = ["学生1", "班级1", "学生2", "班级2", "相似度",
                   "文本相似度", "代码相似度", "风险等级"]

        # 表头
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row, col, header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color=self.color_header, end_color=self.color_header, fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        # 数据行
        for result in detection_result.cross_class_results:
            row += 1

            ws.cell(row, 1, result.student_id)
            ws.cell(row, 2, result.metadata.get('class_name_1', ''))
            ws.cell(row, 3, result.similar_to)
            ws.cell(row, 4, result.metadata.get('class_name_2', ''))
            ws.cell(row, 5, f"{result.overall_similarity:.1f}%")
            ws.cell(row, 6, f"{result.text_similarity:.1f}%")
            ws.cell(row, 7, f"{result.code_similarity:.1f}%")

            # 风险等级
            risk = self._get_risk_level(result.overall_similarity)
            ws.cell(row, 8, risk)

            # 颜色标记
            if result.overall_similarity >= 85:
                for col in range(1, 9):
                    ws.cell(row, col).fill = PatternFill(
                        start_color=self.color_high,
                        end_color=self.color_high,
                        fill_type="solid"
                    )
            elif result.overall_similarity >= 70:
                for col in range(1, 9):
                    ws.cell(row, col).fill = PatternFill(
                        start_color=self.color_medium,
                        end_color=self.color_medium,
                        fill_type="solid"
                    )

        # 调整列宽
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 20
        ws.column_dimensions['E'].width = 12
        ws.column_dimensions['F'].width = 12
        ws.column_dimensions['G'].width = 12
        ws.column_dimensions['H'].width = 10

    def _create_comparison_sheet(self, wb, detection_result):
        """创建班级对比工作表"""
        ws = wb.create_sheet("班级对比")

        row = 1
        headers = ["班级1", "班级2", "平均相似度", "最高相似度",
                   "可疑对数", "平均分差异", "提交率差异"]

        # 表头
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row, col, header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color=self.color_header, end_color=self.color_header, fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        # 数据行
        for comparison in detection_result.class_comparisons:
            row += 1

            ws.cell(row, 1, comparison['class_name_1'])
            ws.cell(row, 2, comparison['class_name_2'])
            ws.cell(row, 3, f"{comparison['avg_similarity']:.1f}%")
            ws.cell(row, 4, f"{comparison['max_similarity']:.1f}%")
            ws.cell(row, 5, comparison['suspicious_pairs'])
            ws.cell(row, 6, f"{comparison['avg_score_diff']:.1f}")
            ws.cell(row, 7, f"{comparison['submission_rate_diff']:.1f}%")

        # 调整列宽
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 12
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 12
        ws.column_dimensions['F'].width = 12
        ws.column_dimensions['G'].width = 12

    def _create_class_detail_sheet(self, wb, class_id: str, class_result):
        """创建班级详细结果工作表"""
        # 使用简短的班级名称作为工作表名称
        sheet_name = class_result.class_name[:10]  # Excel工作表名称限制
        ws = wb.create_sheet(sheet_name)

        row = 1
        headers = ["学生1", "学生2", "相似度", "文本相似度",
                   "代码相似度", "是否跨组", "共享段落数"]

        # 表头
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row, col, header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color=self.color_header, end_color=self.color_header, fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        # 可疑对数据
        for result in class_result.suspicious_results:
            row += 1

            ws.cell(row, 1, result.student_id)
            ws.cell(row, 2, result.similar_to)
            ws.cell(row, 3, f"{result.overall_similarity:.1f}%")
            ws.cell(row, 4, f"{result.text_similarity:.1f}%")
            ws.cell(row, 5, f"{result.code_similarity:.1f}%")
            ws.cell(row, 6, "是" if result.is_cross_group else "否")
            ws.cell(row, 7, len(result.shared_paragraphs))

            # 颜色标记
            if result.overall_similarity >= 85:
                for col in range(1, 8):
                    ws.cell(row, col).fill = PatternFill(
                        start_color=self.color_high,
                        end_color=self.color_high,
                        fill_type="solid"
                    )
            elif result.overall_similarity >= 70:
                for col in range(1, 8):
                    ws.cell(row, col).fill = PatternFill(
                        start_color=self.color_medium,
                        end_color=self.color_medium,
                        fill_type="solid"
                    )

        # 调整列宽
        for col in range(1, 8):
            ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 12

    def _get_risk_level(self, similarity: float) -> str:
        """获取风险等级"""
        if similarity >= 85:
            return "高"
        elif similarity >= 70:
            return "中"
        elif similarity >= 60:
            return "低"
        else:
            return "正常"

    def generate_json(self, detection_result) -> Path:
        """生成JSON报告"""
        output = {
            'meta': {
                'project_name': self.project_name,
                'generated_at': datetime.now().isoformat()
            },
            'summary': detection_result.get_summary(),
            'classes': {},
            'cross_class_results': [],
            'class_comparisons': []
        }

        # 班级结果
        for class_id, result in detection_result.class_results.items():
            output['classes'][class_id] = {
                'class_name': result.class_name,
                'student_count': result.student_count,
                'suspicious_pairs': result.suspicious_pairs,
                'suspicious_rate': result.suspicious_pairs / result.student_count if result.student_count > 0 else 0,
                'group_count': len(result.groups),
                'groups': result.groups
            }

        # 跨班级结果
        for result in detection_result.cross_class_results:
            output['cross_class_results'].append({
                'student1': result.student_id,
                'class1': result.metadata.get('class_name_1', ''),
                'student2': result.similar_to,
                'class2': result.metadata.get('class_name_2', ''),
                'similarity': round(result.overall_similarity, 1),
                'text_similarity': round(result.text_similarity, 1),
                'code_similarity': round(result.code_similarity, 1),
                'risk_level': self._get_risk_level(result.overall_similarity)
            })

        # 班级对比
        output['class_comparisons'] = detection_result.class_comparisons

        # 保存
        filename = f"{self.project_name}_多班级汇总_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        output_path = self.output_dir / filename

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output, f, ensure_ascii=False, indent=2)

        return output_path

    def load_grading_data(self, class_configs: List[Dict]):
        """
        加载各班级的评分数据

        Args:
            class_configs: 班级配置列表，每项包含 class_id, experiment_dir
        """
        for config in class_configs:
            class_id = config['class_id']
            experiment_dir = Path(config.get('experiment_dir', ''))

            # 尝试加载评分结果
            grading_json = experiment_dir / 'results' / 'grading_results.json'

            if grading_json.exists():
                try:
                    with open(grading_json, 'r', encoding='utf-8') as f:
                        data = json.load(f)

                    # 处理两种可能的格式
                    # 格式1: 列表格式 [{student_id, total_score, ...}, ...]
                    # 格式2: 字典格式 {grades: {student_id: {score, grade, ...}}}
                    if isinstance(data, list):
                        self.grading_data[class_id] = data
                    elif isinstance(data, dict):
                        if 'grades' in data and isinstance(data['grades'], dict):
                            # 转换格式2为格式1
                            grading_list = []
                            for student_id, grade_info in data['grades'].items():
                                if isinstance(grade_info, dict):
                                    grading_list.append({
                                        'student_id': student_id,
                                        'total_score': grade_info.get('score', 0),
                                        'grade': grade_info.get('grade', 'F'),
                                        'name': grade_info.get('name', '')
                                    })
                            self.grading_data[class_id] = grading_list
                        else:
                            # 其他字典格式，尝试作为列表使用
                            self.grading_data[class_id] = list(data.values()) if data else []
                    else:
                        self.grading_data[class_id] = []

                    print(f"加载评分数据: {class_id} - {len(self.grading_data[class_id])} 条记录")
                except Exception as e:
                    print(f"警告: 无法加载 {class_id} 的评分数据: {e}")
                    self.grading_data[class_id] = []

    def _create_grading_comparison_sheet(self, wb, detection_result):
        """创建评分对比工作表"""
        ws = wb.create_sheet("评分对比")

        if not self.grading_data:
            # 如果没有评分数据，添加提示信息
            ws['A1'] = "暂无评分数据"
            ws['A1'].font = Font(italic=True, color="808080")
            return

        row = 1
        headers = ["班级ID", "班级名称", "学生数", "平均分", "最高分", "最低分",
                   "及格率", "优秀率", "A等", "B等", "C等", "D等", "F等"]

        # 表头
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row, col, header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color=self.color_header, end_color=self.color_header, fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        # 收集各班级的评分统计
        class_stats = []

        for class_id in detection_result.class_results.keys():
            if class_id not in self.grading_data:
                continue

            grading_list = self.grading_data[class_id]

            if not grading_list:
                continue

            # 过滤掉非字典类型的数据
            grading_list = [g for g in grading_list if isinstance(g, dict)]

            if not grading_list:
                continue

            # 计算统计数据
            scores = [g.get('total_score', 0) for g in grading_list if isinstance(g.get('total_score'), (int, float))]
            grades = [g.get('grade', 'F') for g in grading_list if isinstance(g, dict)]

            if not scores:
                continue

            avg_score = sum(scores) / len(scores)
            max_score = max(scores)
            min_score = min(scores)

            # 计算及格率和优秀率
            passing = sum(1 for s in scores if s >= 60)
            excellent = sum(1 for s in scores if s >= 90)
            passing_rate = passing / len(scores) * 100 if scores else 0
            excellent_rate = excellent / len(scores) * 100 if scores else 0

            # 统计各等级人数
            grade_counts = {
                'A': grades.count('A'),
                'B': grades.count('B'),
                'C': grades.count('C'),
                'D': grades.count('D'),
                'F': grades.count('F')
            }

            # 获取班级名称
            class_name = class_id
            if class_id in detection_result.class_results:
                class_name = detection_result.class_results[class_id].class_name

            class_stats.append({
                'class_id': class_id,
                'class_name': class_name,
                'student_count': len(scores),
                'avg_score': round(avg_score, 1),
                'max_score': round(max_score, 1),
                'min_score': round(min_score, 1),
                'passing_rate': round(passing_rate, 1),
                'excellent_rate': round(excellent_rate, 1),
                'grade_counts': grade_counts
            })

        # 数据行
        for stats in class_stats:
            row += 1
            ws.cell(row, 1, stats['class_id'])
            ws.cell(row, 2, stats['class_name'])
            ws.cell(row, 3, stats['student_count'])
            ws.cell(row, 4, stats['avg_score'])
            ws.cell(row, 5, stats['max_score'])
            ws.cell(row, 6, stats['min_score'])
            ws.cell(row, 7, f"{stats['passing_rate']:.1f}%")
            ws.cell(row, 8, f"{stats['excellent_rate']:.1f}%")
            ws.cell(row, 9, stats['grade_counts']['A'])
            ws.cell(row, 10, stats['grade_counts']['B'])
            ws.cell(row, 11, stats['grade_counts']['C'])
            ws.cell(row, 12, stats['grade_counts']['D'])
            ws.cell(row, 13, stats['grade_counts']['F'])

        # 调整列宽
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 15
        for col in range(3, 14):
            ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 10

    def _generate_comparison_charts(self, detection_result) -> Optional[Path]:
        """
        生成班级对比图表

        Returns:
            图像文件路径或None
        """
        if not HAS_MATPLOTLIB or not self.grading_data:
            return None

        # 创建图表
        fig, axes = plt.subplots(2, 2, figsize=(14, 10))
        fig.suptitle(f'{self.project_name} - 班级对比分析', fontsize=16, fontweight='bold')

        # 收集数据
        class_names = []
        avg_scores = []
        passing_rates = []
        excellent_rates = []
        suspicious_rates = []

        for class_id, result in detection_result.class_results.items():
            if class_id not in self.grading_data:
                continue

            grading_list = self.grading_data[class_id]
            if not grading_list:
                continue

            class_names.append(result.class_name)

            # 计算平均分
            scores = [g.get('total_score', 0) for g in grading_list if g.get('total_score') is not None]
            avg_scores.append(sum(scores) / len(scores) if scores else 0)

            # 计算及格率和优秀率
            passing = sum(1 for s in scores if s >= 60)
            excellent = sum(1 for s in scores if s >= 90)
            passing_rates.append(passing / len(scores) * 100 if scores else 0)
            excellent_rates.append(excellent / len(scores) * 100 if scores else 0)

            # 可疑率
            suspicious_rate = result.suspicious_pairs / result.student_count if result.student_count > 0 else 0
            suspicious_rates.append(suspicious_rate * 100)

        if not class_names:
            plt.close(fig)
            return None

        # 1. 平均分对比（柱状图）
        ax1 = axes[0, 0]
        colors1 = ['#4CAF50' if s >= 70 else '#FF9800' if s >= 60 else '#F44336' for s in avg_scores]
        ax1.bar(class_names, avg_scores, color=colors1)
        ax1.set_title('班级平均分对比')
        ax1.set_ylabel('平均分')
        ax1.set_ylim(0, 100)
        for i, v in enumerate(avg_scores):
            ax1.text(i, v + 1, f'{v:.1f}', ha='center', va='bottom')

        # 2. 及格率对比（柱状图）
        ax2 = axes[0, 1]
        ax2.bar(class_names, passing_rates, color='#2196F3')
        ax2.set_title('及格率对比')
        ax2.set_ylabel('及格率 (%)')
        ax2.set_ylim(0, 100)
        for i, v in enumerate(passing_rates):
            ax2.text(i, v + 1, f'{v:.1f}%', ha='center', va='bottom')

        # 3. 优秀率对比（柱状图）
        ax3 = axes[1, 0]
        ax3.bar(class_names, excellent_rates, color='#FFC107')
        ax3.set_title('优秀率对比 (≥90分)')
        ax3.set_ylabel('优秀率 (%)')
        ax3.set_ylim(0, 100)
        for i, v in enumerate(excellent_rates):
            ax3.text(i, v + 1, f'{v:.1f}%', ha='center', va='bottom')

        # 4. 可疑率对比（柱状图）
        ax4 = axes[1, 1]
        ax4.bar(class_names, suspicious_rates, color='#FF5722')
        ax4.set_title('可疑率对比')
        ax4.set_ylabel('可疑率 (%)')
        for i, v in enumerate(suspicious_rates):
            ax4.text(i, v + 1, f'{v:.1f}%', ha='center', va='bottom')

        plt.tight_layout()

        # 保存图表
        chart_path = self.output_dir / f"{self.project_name}_对比图表_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        plt.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close(fig)

        return chart_path

    def generate_pdf(self, detection_result) -> Optional[Path]:
        """生成PDF报告"""
        if not HAS_REPORTLAB:
            return None

        # 创建PDF文档
        filename = f"{self.project_name}_多班级汇总_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        output_path = self.output_dir / filename

        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )

        # 创建内容列表
        story = []
        styles = getSampleStyleSheet()

        # 添加自定义样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#2C3E50'),
            spaceAfter=12,
            alignment=TA_CENTER
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#34495E'),
            spaceAfter=10
        )

        # 标题
        story.append(Paragraph(f"{self.project_name} - 多班级查重汇总", title_style))
        story.append(Paragraph(f"生成时间: {detection_result.timestamp}", styles['Normal']))
        story.append(Spacer(1, 0.5*cm))

        # 汇总统计
        story.append(Paragraph("汇总统计", heading_style))
        summary = detection_result.get_summary()

        summary_data = [
            ["统计项", "数值"],
            ["班级数量", str(summary['total_classes'])],
            ["总学生数", str(summary['total_students'])],
            ["班级内可疑对", str(summary['total_suspicious_pairs'])],
            ["跨班级可疑对", str(summary['cross_class_suspicious_pairs'])],
            ["班级对比数", str(summary['class_comparisons'])]
        ]

        summary_table = Table(summary_data, colWidths=[8*cm, 4*cm])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (1, 0), colors.HexColor('#3498DB')),
            ('TEXTCOLOR', (0, 0), (1, 0), colors.white),
            ('ALIGN', (0, 0), (1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (1, -1), 6),
            ('GRID', (0, 0), (1, -1), 1, colors.black)
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.5*cm))

        # 班级概览
        story.append(Paragraph("班级概览", heading_style))

        class_overview_data = [["班级ID", "班级名称", "学生数", "可疑对数", "可疑率", "团伙数"]]

        for class_id, result in detection_result.class_results.items():
            suspicious_rate = result.suspicious_pairs / result.student_count if result.student_count > 0 else 0
            class_overview_data.append([
                class_id,
                result.class_name,
                str(result.student_count),
                str(result.suspicious_pairs),
                f"{suspicious_rate:.1%}",
                str(len(result.groups))
            ])

        class_table = Table(class_overview_data, colWidths=[4*cm, 3*cm, 2*cm, 2*cm, 2*cm, 2*cm])
        class_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (5, 0), colors.HexColor('#3498DB')),
            ('TEXTCOLOR', (0, 0), (5, 0), colors.white),
            ('ALIGN', (0, 0), (5, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (5, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (5, -1), 9),
            ('BOTTOMPADDING', (0, 0), (5, -1), 6),
            ('GRID', (0, 0), (5, -1), 1, colors.black)
        ]))
        story.append(class_table)
        story.append(Spacer(1, 0.5*cm))

        # 班级对比
        if detection_result.class_comparisons:
            story.append(Paragraph("班级对比", heading_style))

            comparison_data = [["班级1", "班级2", "平均相似度", "最高相似度", "可疑对数"]]

            for comp in detection_result.class_comparisons:
                comparison_data.append([
                    comp['class_name_1'],
                    comp['class_name_2'],
                    f"{comp['avg_similarity']:.1f}%",
                    f"{comp['max_similarity']:.1f}%",
                    str(comp['suspicious_pairs'])
                ])

            comparison_table = Table(comparison_data, colWidths=[3*cm, 3*cm, 2*cm, 2*cm, 2*cm])
            comparison_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (4, 0), colors.HexColor('#3498DB')),
                ('TEXTCOLOR', (0, 0), (4, 0), colors.white),
                ('ALIGN', (0, 0), (4, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (4, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (4, -1), 9),
                ('BOTTOMPADDING', (0, 0), (4, -1), 6),
                ('GRID', (0, 0), (4, -1), 1, colors.black)
            ]))
            story.append(comparison_table)
            story.append(Spacer(1, 0.5*cm))

        # 评分对比
        if self.grading_data:
            story.append(PageBreak())
            story.append(Paragraph("评分对比", heading_style))

            grading_data = [["班级ID", "班级名称", "学生数", "平均分", "及格率", "优秀率"]]

            for class_id in detection_result.class_results.keys():
                if class_id not in self.grading_data:
                    continue

                grading_list = self.grading_data[class_id]
                grading_list = [g for g in grading_list if isinstance(g, dict)]

                if not grading_list:
                    continue

                scores = [g.get('total_score', 0) for g in grading_list if isinstance(g.get('total_score'), (int, float))]

                if scores:
                    avg_score = sum(scores) / len(scores)
                    passing = sum(1 for s in scores if s >= 60)
                    excellent = sum(1 for s in scores if s >= 90)
                    passing_rate = passing / len(scores) * 100
                    excellent_rate = excellent / len(scores) * 100

                    result = detection_result.class_results[class_id]
                    grading_data.append([
                        class_id,
                        result.class_name,
                        str(len(scores)),
                        f"{avg_score:.1f}",
                        f"{passing_rate:.1f}%",
                        f"{excellent_rate:.1f}%"
                    ])

            grading_table = Table(grading_data, colWidths=[4*cm, 3*cm, 2*cm, 2*cm, 2*cm, 2*cm])
            grading_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (5, 0), colors.HexColor('#3498DB')),
                ('TEXTCOLOR', (0, 0), (5, 0), colors.white),
                ('ALIGN', (0, 0), (5, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (5, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (5, -1), 9),
                ('BOTTOMPADDING', (0, 0), (5, -1), 6),
                ('GRID', (0, 0), (5, -1), 1, colors.black)
            ]))
            story.append(grading_table)

        # 构建PDF
        try:
            doc.build(story)
            return output_path
        except Exception:
            logger.exception("PDF 构建失败: %s", output_path)
            return None

    def generate_word(self, detection_result) -> Optional[Path]:
        """生成Word报告"""
        if not HAS_PYTHON_DOCX:
            return None

        # 创建Word文档
        filename = f"{self.project_name}_多班级汇总_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = self.output_dir / filename

        doc = Document()

        # 设置默认字体
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # 标题
        title = doc.add_heading(f"{self.project_name} - 多班级查重汇总", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 生成时间
        doc.add_paragraph(f"生成时间: {detection_result.timestamp}")

        # 汇总统计
        doc.add_heading("汇总统计", 1)
        summary = detection_result.get_summary()

        # 创建汇总表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "统计项"
        hdr_cells[1].text = "数值"

        summary_items = [
            ("班级数量", str(summary['total_classes'])),
            ("总学生数", str(summary['total_students'])),
            ("班级内可疑对", str(summary['total_suspicious_pairs'])),
            ("跨班级可疑对", str(summary['cross_class_suspicious_pairs'])),
            ("班级对比数", str(summary['class_comparisons']))
        ]

        for label, value in summary_items:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        # 班级概览
        doc.add_heading("班级概览", 1)

        class_table = doc.add_table(rows=1, cols=6)
        class_table.style = 'Light Grid Accent 1'

        hdr_cells = class_table.rows[0].cells
        headers = ["班级ID", "班级名称", "学生数", "可疑对数", "可疑率", "团伙数"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        for class_id, result in detection_result.class_results.items():
            row_cells = class_table.add_row().cells
            suspicious_rate = result.suspicious_pairs / result.student_count if result.student_count > 0 else 0

            row_cells[0].text = class_id
            row_cells[1].text = result.class_name
            row_cells[2].text = str(result.student_count)
            row_cells[3].text = str(result.suspicious_pairs)
            row_cells[4].text = f"{suspicious_rate:.1%}"
            row_cells[5].text = str(len(result.groups))

        # 班级对比
        if detection_result.class_comparisons:
            doc.add_heading("班级对比", 1)

            comp_table = doc.add_table(rows=1, cols=5)
            comp_table.style = 'Light Grid Accent 1'

            hdr_cells = comp_table.rows[0].cells
            comp_headers = ["班级1", "班级2", "平均相似度", "最高相似度", "可疑对数"]
            for i, header in enumerate(comp_headers):
                hdr_cells[i].text = header

            for comp in detection_result.class_comparisons:
                row_cells = comp_table.add_row().cells
                row_cells[0].text = comp['class_name_1']
                row_cells[1].text = comp['class_name_2']
                row_cells[2].text = f"{comp['avg_similarity']:.1f}%"
                row_cells[3].text = f"{comp['max_similarity']:.1f}%"
                row_cells[4].text = str(comp['suspicious_pairs'])

        # 评分对比
        if self.grading_data:
            doc.add_page_break()
            doc.add_heading("评分对比", 1)

            grading_table = doc.add_table(rows=1, cols=6)
            grading_table.style = 'Light Grid Accent 1'

            hdr_cells = grading_table.rows[0].cells
            grading_headers = ["班级ID", "班级名称", "学生数", "平均分", "及格率", "优秀率"]
            for i, header in enumerate(grading_headers):
                hdr_cells[i].text = header

            for class_id in detection_result.class_results.keys():
                if class_id not in self.grading_data:
                    continue

                grading_list = self.grading_data[class_id]
                grading_list = [g for g in grading_list if isinstance(g, dict)]

                if not grading_list:
                    continue

                scores = [g.get('total_score', 0) for g in grading_list if isinstance(g.get('total_score'), (int, float))]

                if scores:
                    avg_score = sum(scores) / len(scores)
                    passing = sum(1 for s in scores if s >= 60)
                    excellent = sum(1 for s in scores if s >= 90)
                    passing_rate = passing / len(scores) * 100
                    excellent_rate = excellent / len(scores) * 100

                    result = detection_result.class_results[class_id]
                    row_cells = grading_table.add_row().cells
                    row_cells[0].text = class_id
                    row_cells[1].text = result.class_name
                    row_cells[2].text = str(len(scores))
                    row_cells[3].text = f"{avg_score:.1f}"
                    row_cells[4].text = f"{passing_rate:.1f}%"
                    row_cells[5].text = f"{excellent_rate:.1f}%"

        # 保存文档
        try:
            doc.save(output_path)
            return output_path
        except Exception as e:
            print(f"Word保存失败: {e}")
            return None

