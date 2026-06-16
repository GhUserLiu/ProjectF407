#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
图像质量检测模块
Image Quality Checker

检测报告中的图片质量、相关性、符合度
"""

import re
import io
from dataclasses import dataclass, field
from typing import List, Dict, Tuple, Optional
from enum import Enum
from pathlib import Path

# Lazy import PIL and NumPy to avoid PyInstaller conflicts
# They will be imported only when actually needed
PILLOW_AVAILABLE = True  # Assume available, will verify on first use
_pil_Image = None
_np = None

def _get_pil_modules():
    """Lazy import PIL and NumPy to avoid initialization conflicts"""
    global _pil_Image, _np, PILLOW_AVAILABLE
    if _pil_Image is None:
        try:
            from PIL import Image as _PIL_Image
            import numpy as _numpy
            _pil_Image = _PIL_Image
            _np = _numpy
            PILLOW_AVAILABLE = True
        except ImportError:
            PILLOW_AVAILABLE = False
            _pil_Image = False
            _np = False
    return _pil_Image, _np


class ImageQuality(Enum):
    """图像质量等级"""
    EXCELLENT = "excellent"  # 优秀 - 清晰、相关、规范
    GOOD = "good"           # 良好 - 基本符合要求
    ACCEPTABLE = "acceptable"  # 可接受 - 有小问题
    POOR = "poor"           # 差 - 存在明显问题
    CRITICAL = "critical"   # 严重 - 需要重新拍摄


@dataclass
class ImageIssue:
    """图像问题"""
    severity: str
    category: str
    message: str
    suggestion: str = ""


@dataclass
class ImageAnalysisResult:
    """图像分析结果"""
    image_count: int
    total_score: float      # 总分 0-100
    quality_rating: ImageQuality
    issues: List[ImageIssue]
    strengths: List[str]
    details: List[Dict]     # 每张图片的详细信息


class ImageQualityChecker:
    """图像质量检查器"""

    # 实验相关的关键词
    EXPERIMENT_KEYWORDS = {
        "档位实验": [
            "LED", "档位", "电路", "连接", "接线",
            "开发板", "STM32", "按键", "消抖",
            "PF9", "PF10", "PE4", "GPIO"
        ],
        "转向灯实验": [
            "LED", "转向灯", "闪烁", "电路",
            "开发板", "STM32", "延时", "GPIO"
        ]
    }

    def __init__(self, experiment_type: str = "档位实验"):
        """
        初始化图像质量检查器

        Args:
            experiment_type: 实验类型
        """
        self.experiment_type = experiment_type
        self.keywords = self.EXPERIMENT_KEYWORDS.get(experiment_type, [])

    def check_report_images(
        self,
        docx_path: Path,
        experiment_type: str = "档位实验"
    ) -> ImageAnalysisResult:
        """
        检查报告中的图片质量

        Args:
            docx_path: 报告文件路径
            experiment_type: 实验类型

        Returns:
            图像分析结果
        """
        # Get PIL modules lazily
        Image, np = _get_pil_modules()
        if not PILLOW_AVAILABLE or Image is None:
            return ImageAnalysisResult(
                image_count=0,
                total_score=50,
                quality_rating=ImageQuality.ACCEPTABLE,
                issues=[
                    ImageIssue(
                        severity="medium",
                        category="依赖缺失",
                        message="Pillow库未安装，无法进行详细图像分析",
                        suggestion="安装Pillow: pip install Pillow"
                    )
                ],
                strengths=[],
                details=[]
            )

        try:
            from docx import Document

            doc = Document(docx_path)
            images = []
            issues = []
            strengths = []
            details = []

            # 提取图片
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        img = Image.open(io.BytesIO(image_data))  # Image from lazy import
                        images.append(img)
                    except Exception as e:
                        issues.append(ImageIssue(
                            severity="low",
                            category="图片提取",
                            message=f"无法读取某张图片: {str(e)}",
                            suggestion=""
                        ))

            if not images:
                return ImageAnalysisResult(
                    image_count=0,
                    total_score=0,
                    quality_rating=ImageQuality.CRITICAL,
                    issues=[
                        ImageIssue(
                            severity="critical",
                            category="图片缺失",
                            message="报告中未检测到任何图片",
                            suggestion="请添加实验结果照片或截图"
                        )
                    ],
                    strengths=[],
                    details=[]
                )

            # 检查每张图片
            total_score = 0
            for i, img in enumerate(images):
                result = self._check_single_image(img, i + 1)
                details.append(result)
                total_score += result['score']

                # 收集问题和亮点
                for issue in result['issues']:
                    issues.append(issue)
                for strength in result['strengths']:
                    strengths.append(strength)

            # 计算平均分
            avg_score = total_score / len(images) if images else 0

            # 确定质量等级
            if avg_score >= 85:
                quality_rating = ImageQuality.EXCELLENT
            elif avg_score >= 70:
                quality_rating = ImageQuality.GOOD
            elif avg_score >= 50:
                quality_rating = ImageQuality.ACCEPTABLE
            elif avg_score >= 30:
                quality_rating = ImageQuality.POOR
            else:
                quality_rating = ImageQuality.CRITICAL

            return ImageAnalysisResult(
                image_count=len(images),
                total_score=round(avg_score, 1),
                quality_rating=quality_rating,
                issues=issues,
                strengths=strengths,
                details=details
            )

        except Exception as e:
            return ImageAnalysisResult(
                image_count=0,
                total_score=0,
                quality_rating=ImageQuality.CRITICAL,
                issues=[
                    ImageIssue(
                        severity="critical",
                        category="文件读取",
                        message=f"无法读取报告文件: {str(e)}",
                        suggestion="请确认文件格式正确"
                    )
                ],
                strengths=[],
                details=[]
            )

    def _check_single_image(self, img, index: int) -> Dict:
        """
        检查单张图片质量

        Args:
            img: PIL图片对象 (lazy imported Image.Image)
            index: 图片索引

        Returns:
            图片检查结果
        """
        issues = []
        strengths = []
        score = 100.0

        # 1. 检查分辨率
        width, height = img.size
        if width < 400 or height < 300:
            issues.append(ImageIssue(
                severity="high",
                category="分辨率",
                message=f"图片{index}分辨率过低 ({width}x{height})",
                suggestion="建议使用至少800x600像素的图片"
            ))
            score -= 20
        elif width < 800 or height < 600:
            issues.append(ImageIssue(
                severity="medium",
                category="分辨率",
                message=f"图片{index}分辨率偏低 ({width}x{height})",
                suggestion="建议使用更高分辨率的图片"
            ))
            score -= 10
        else:
            strengths.append(f"图片{index}分辨率足够")

        # 2. 检查清晰度（基于拉普拉斯方差）
        try:
            # 转换为灰度图
            gray_img = img.convert('L')
            # Get np from lazy import
            _, np = _get_pil_modules()
            if PILLOW_AVAILABLE and np is not None:
                img_array = np.array(gray_img)
                laplacian_var = np.var(img_array)

                if laplacian_var < 100:
                    issues.append(ImageIssue(
                        severity="high",
                        category="清晰度",
                        message=f"图片{index}可能模糊",
                        suggestion="建议使用更清晰的照片"
                    ))
                    score -= 15
                elif laplacian_var > 500:
                    strengths.append(f"图片{index}清晰度高")
        except:
            pass

        # 3. 检查图片格式
        if img.format not in ['JPEG', 'PNG', 'JPG']:
            issues.append(ImageIssue(
                severity="low",
                category="格式",
                message=f"图片{index}格式为 {img.format or '未知'}",
                suggestion="建议使用JPEG或PNG格式"
            ))
            score -= 5

        # 4. 检查图片大小
        try:
            img_bytes = io.BytesIO()
            img.save(img_bytes, format='JPEG' if img.format != 'PNG' else 'PNG')
            size_kb = len(img_bytes.getvalue()) / 1024

            if size_kb > 2000:
                issues.append(ImageIssue(
                    severity="low",
                    category="文件大小",
                    message=f"图片{index}文件过大 ({size_kb:.0f}KB)",
                    suggestion="建议压缩图片以减小报告文件大小"
                ))
                score -= 5
            elif size_kb < 10:
                issues.append(ImageIssue(
                    severity="medium",
                    category="文件大小",
                    message=f"图片{index}文件过小，可能是缩略图 ({size_kb:.0f}KB)",
                    suggestion="请使用原始大小的图片"
                ))
                score -= 10
        except:
            pass

        # 5. 检查宽高比（截图检查）
        aspect_ratio = width / height if height > 0 else 1
        if 0.8 <= aspect_ratio <= 1.3:
            # 接近正方形，可能是电路图
            strengths.append(f"图片{index}可能是电路图或设备照片")
        elif aspect_ratio > 1.5:
            # 宽图片，可能是代码截图
            if 'code' not in [i.category for i in issues]:
                issues.append(ImageIssue(
                    severity="info",
                    category="内容类型",
                    message=f"图片{index}可能是代码截图",
                    suggestion="代码图片建议使用等宽字体格式，不要使用截图"
                ))

        return {
            'index': index,
            'size': (width, height),
            'format': img.format,
            'score': max(0, score),
            'issues': issues,
            'strengths': strengths
        }

    def check_text_image_references(self, report_text: str) -> List[ImageIssue]:
        """
        检查文本中对图片的引用

        Args:
            report_text: 报告文本

        Returns:
            问题列表
        """
        issues = []

        # 检查是否有图片引用
        image_refs = re.findall(r'[图图]\s*\d+|Figure\s*\d+|图片|截图|照片', report_text)

        if not image_refs:
            issues.append(ImageIssue(
                severity="high",
                category="图片引用",
                message="报告中没有对图片的引用",
                suggestion="建议在正文中引用图片，如'如图1所示'"
            ))
        else:
            # 检查引用格式
            invalid_refs = re.findall(r'[图图]\s*[^\d]|图片[^\d]', report_text)
            if invalid_refs:
                issues.append(ImageIssue(
                    severity="low",
                    category="图片引用",
                    message="图片引用格式不规范",
                    suggestion="使用标准格式：'如图1所示'、'图1-1'等"
                ))

        return issues

    def generate_image_quality_report(self, result: ImageAnalysisResult) -> str:
        """
        生成图片质量报告

        Args:
            result: 图像分析结果

        Returns:
            Markdown格式报告
        """
        lines = [
            "# 图片质量分析报告",
            "",
            f"**检测图片数量**: {result.image_count}",
            f"**平均得分**: {result.total_score}/100",
            f"**质量等级**: {result.quality_rating.value}",
            ""
        ]

        # 亮点
        if result.strengths:
            lines.append("## ✅ 优点")
            lines.append("")
            for strength in result.strengths:
                lines.append(f"- {strength}")
            lines.append("")

        # 问题
        if result.issues:
            severity_order = ["critical", "high", "medium", "low", "info"]
            grouped = {}
            for issue in result.issues:
                if issue.severity not in grouped:
                    grouped[issue.severity] = []
                grouped[issue.severity].append(issue)

            lines.append("## ⚠️ 问题与建议")
            lines.append("")

            for severity in severity_order:
                if severity in grouped:
                    emoji = {
                        "critical": "🔴",
                        "high": "🟠",
                        "medium": "🟡",
                        "low": "🟢",
                        "info": "🔵"
                    }
                    lines.append(f"### {emoji[severity]} {severity.upper()} 级")
                    lines.append("")
                    for issue in grouped[severity]:
                        lines.append(f"- **{issue.category}**: {issue.message}")
                        if issue.suggestion:
                            lines.append(f"  - 建议: {issue.suggestion}")
                    lines.append("")

        # 详细信息
        if result.details:
            lines.append("## 📊 详细信息")
            lines.append("")
            for detail in result.details:
                lines.append(f"### 图片 {detail['index']}")
                lines.append(f"- 尺寸: {detail['size'][0]}x{detail['size'][1]}")
                lines.append(f"- 格式: {detail['format']}")
                lines.append(f"- 得分: {detail['score']:.1f}")
                lines.append("")

        return '\n'.join(lines)


class ImageRelevanceChecker:
    """图片相关性检查器"""

    @staticmethod
    def check_image_relevance(
        docx_path: Path,
        experiment_type: str = "档位实验"
    ) -> Tuple[float, List[str]]:
        """
        检查图片与实验的相关性

        Args:
            docx_path: 报告文件路径
            experiment_type: 实验类型

        Returns:
            (相关性得分, 问题列表)
        """
        # 获取实验关键词
        keywords = ImageQualityChecker.EXPERIMENT_KEYWORDS.get(
            experiment_type,
            ["实验", "结果", "电路", "代码"]
        )

        issues = []
        relevance_score = 100.0

        try:
            from docx import Document

            doc = Document(docx_path)

            # 提取所有文本
            full_text = '\n'.join([p.text for p in doc.paragraphs])

            # 检查图片数量
            image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.target_ref)

            if image_count == 0:
                issues.append("报告中没有图片")
                relevance_score = 0
            elif image_count < 2:
                issues.append("图片数量偏少，建议至少添加2-3张图片")
                relevance_score -= 30
            elif image_count >= 5:
                issues.append("图片数量较多，注意与内容的相关性")
                # 图片多不一定是坏事，不减分

            # 检查图片引用
            image_refs = len(re.findall(r'[图图]\s*\d+|Figure\s*\d+', full_text))
            if image_refs == 0 and image_count > 0:
                issues.append("添加了图片但未在正文中引用")
                relevance_score -= 20

            # 检查关键场景是否有图片
            critical_scenes = {
                "档位实验": ["电路", "接线", "硬件", "LED", "现象"],
                "转向灯实验": ["电路", "LED", "现象", "闪烁"]
            }

            required = critical_scenes.get(experiment_type, [])

            # 简化检查：看文本是否提到这些场景
            for scene in required:
                if scene in full_text:
                    # 文本提到了，检查是否有图片支持
                    if image_count < len(required):
                        issues.append(f"提到了{scene}，但可能缺少相应图片")
                        relevance_score -= 10

            return max(0, relevance_score), issues

        except Exception as e:
            return 0, [f"无法分析图片相关性: {str(e)}"]


def check_images_from_directory(
    image_dir: Path,
    experiment_type: str = "档位实验"
) -> ImageAnalysisResult:
    """
    从目录检查图片质量

    Args:
        image_dir: 图片目录
        experiment_type: 实验类型

    Returns:
        图片分析结果
    """
    # Get PIL modules lazily
    Image, np = _get_pil_modules()
    if not PILLOW_AVAILABLE or Image is None:
        return ImageAnalysisResult(
            image_count=0,
            total_score=0,
            quality_rating=ImageQuality.CRITICAL,
            issues=[
                ImageIssue(
                    severity="critical",
                    category="依赖缺失",
                    message="Pillow库未安装",
                    suggestion="安装Pillow: pip install Pillow"
                )
            ],
            strengths=[],
            details=[]
        )

    checker = ImageQualityChecker(experiment_type)

    # 获取所有图片文件
    image_files = list(image_dir.glob("*.jpg")) + list(image_dir.glob("*.png")) + \
                  list(image_dir.glob("*.jpeg"))

    if not image_files:
        return ImageAnalysisResult(
            image_count=0,
            total_score=0,
            quality_rating=ImageQuality.CRITICAL,
            issues=[
                ImageIssue(
                    severity="critical",
                    category="图片缺失",
                    message=f"目录中没有图片文件: {image_dir}",
                    suggestion="请添加实验结果图片"
                )
            ],
            strengths=[],
            details=[]
        )

    # 检查每张图片
    total_score = 0
    all_issues = []
    all_strengths = []
    details = []

    for i, img_file in enumerate(image_files, 1):
        try:
            img = Image.open(img_file)  # Image from lazy import
            result = checker._check_single_image(img, i)
            result['filename'] = img_file.name
            details.append(result)
            total_score += result['score']

            all_issues.extend(result['issues'])
            all_strengths.extend(result['strengths'])
        except Exception as e:
            all_issues.append(ImageIssue(
                severity="low",
                category="文件读取",
                message=f"无法读取图片 {img_file.name}: {str(e)}",
                suggestion=""
            ))

    avg_score = total_score / len(image_files) if image_files else 0

    # 确定质量等级
    if avg_score >= 85:
        quality_rating = ImageQuality.EXCELLENT
    elif avg_score >= 70:
        quality_rating = ImageQuality.GOOD
    elif avg_score >= 50:
        quality_rating = ImageQuality.ACCEPTABLE
    elif avg_score >= 30:
        quality_rating = ImageQuality.POOR
    else:
        quality_rating = ImageQuality.CRITICAL

    return ImageAnalysisResult(
        image_count=len(image_files),
        total_score=round(avg_score, 1),
        quality_rating=quality_rating,
        issues=all_issues,
        strengths=all_strengths,
        details=details
    )
