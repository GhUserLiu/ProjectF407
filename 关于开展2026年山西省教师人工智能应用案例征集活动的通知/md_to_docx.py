#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Markdown转Word文档"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

# 读取Markdown文件
with open('创AI案例-开发与应用报告.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

# 创建Word文档
doc = Document()

# 设置默认样式
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 辅助函数
def set_chinese_font(run):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 解析Markdown
i = 0
in_code_block = False
code_lines = []
skip_first_title = True  # 跳过第一个标题（已手动添加封面）

while i < len(lines):
    line = lines[i].rstrip('\n\r')

    # 代码块
    if line.startswith('```'):
        if not in_code_block:
            in_code_block = True
            code_lines = []
        else:
            # 代码块结束
            if code_lines:
                para = doc.add_paragraph()
                run = para.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            code_lines = []
            in_code_block = False
        i += 1
        continue

    if in_code_block:
        code_lines.append(line)
        i += 1
        continue

    # 空行
    if not line.strip():
        doc.add_paragraph()
        i += 1
        continue

    # 标题
    if line.startswith('#'):
        match = re.match(r'^(#+)\s+(.+)$', line)
        if match:
            level = len(match.group(1))
            text = match.group(2)

            # 跳过第一个大标题
            if skip_first_title and level == 1:
                skip_first_title = False
                i += 1
                continue

            if level == 1:
                doc.add_heading(text, level=1)
            elif level == 2:
                doc.add_heading(text, level=2)
            elif level == 3:
                doc.add_heading(text, level=3)
            else:
                doc.add_heading(text, level=4)
            i += 1
            continue

    # 表格
    if line.startswith('|') and '|' in line:
        # 检查是否是表格
        if i + 1 < len(lines) and '---' in lines[i + 1]:
            headers = [h.strip() for h in line.split('|')[1:-1] if h.strip()]
            i += 2  # 跳过表头和分隔线

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid'

            # 设置表头
            for j, header in enumerate(headers):
                table.rows[0].cells[j].text = header
                for run in table.rows[0].cells[j].paragraphs[0].runs:
                    run.font.bold = True
                    set_chinese_font(run)

            # 读取表格行
            row_idx = 1
            while i < len(lines) and lines[i].startswith('|'):
                cells = [c.strip() for c in lines[i].split('|')[1:-1] if c.strip()]
                if cells:
                    row = table.add_row()
                    for j, cell_text in enumerate(cells):
                        if j < len(row.cells):
                            row.cells[j].text = cell_text
                            for run in row.cells[j].paragraphs[0].runs:
                                set_chinese_font(run)
                row_idx += 1
                i += 1

            # 添加空行
            doc.add_paragraph()
            continue

    # 列表
    if line.strip().startswith('- '):
        while i < len(lines) and lines[i].strip().startswith('- '):
            item = lines[i].strip()[2:]
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(item)
            set_chinese_font(run)
            i += 1
        continue

    if re.match(r'^\d+\.\s', line.strip()):
        while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
            item = re.sub(r'^\d+\.\s', '', lines[i].strip())
            para = doc.add_paragraph(style='List Number')
            run = para.add_run(item)
            set_chinese_font(run)
            i += 1
        continue

    # 分隔线
    if line.strip() in ['---', '***', '====']:
        doc.add_paragraph('_' * 60)
        i += 1
        continue

    # 普通段落 - 处理内联格式
    para = doc.add_paragraph()
    text = line

    # 处理粗体 **text**
    for match in re.finditer(r'\*\*([^*]+)\*\*', text):
        before = text[:match.start()]
        bold_text = match.group(1)
        after = text[match.end():]

        if before:
            run = para.add_run(before)
            set_chinese_font(run)

        run = para.add_run(bold_text)
        run.font.bold = True
        set_chinese_font(run)

        text = after

    if text:
        run = para.add_run(text)
        set_chinese_font(run)

    i += 1

# 保存
doc.save('创AI案例-开发与应用报告.docx')
print('Word文档生成成功: 创AI案例-开发与应用报告.docx')
