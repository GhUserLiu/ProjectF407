"""
生成输出文件：
1. 教师用Excel工作簿
2. 学生用反馈文档
"""

import json
from pathlib import Path
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).parent.parent.parent.parent.parent  # Go up to project root
# 默认使用最新的实验目录，可通过命令行参数覆盖
EXPERIMENT_DIR = BASE_DIR / "docs" / "teaching" / "2026-春季" / "汽服2302B班" / "07-car-gear"
PROCESSED_DIR = EXPERIMENT_DIR / "processed"
DATA_DIR = Path(__file__).parent.parent / "rubrics"
OUTPUT_DIR = EXPERIMENT_DIR
TEACHER_OUTPUT = OUTPUT_DIR / "results"
STUDENT_OUTPUT = OUTPUT_DIR / "feedback"

def create_teacher_excel(evaluations, rubric, quality_data=None):
    """Create Excel workbook for teacher"""
    print("Creating teacher Excel workbook...")

    # Original plagiarism student list (from auto_score.py)
    ORIGINAL_PLAGIARISM_STUDENTS = {
        '23071140217', '23071140216', '23071140214', '23071140228',
        '23071140233', '23071140220', '23071140223', '23071140213',
        '23071140219', '23071140204', '23071140208'
    }

    wb = Workbook()

    # Header styles
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    warning_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # ============================================================
    # Sheet 1: 总评表（汇总所有实验，目前仅有一个实验）
    # ============================================================
    ws_overall = wb.active
    ws_overall.title = "总评表"

    overall_headers = ['学号', '姓名', '实验07-档位', '实验-未完成', '总分', '等级', '备注']
    ws_overall.append(overall_headers)

    for cell in ws_overall[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = border

    for eval in evaluations:
        student_id = eval['student_id']
        total_score = eval.get('total_score', 0)
        note = []
        grade_label = eval.get('grade_label', '')

        # Check if plagiarist - set score to 0
        if student_id in ORIGINAL_PLAGIARISM_STUDENTS:
            total_score = 0
            grade_label = '不及格'
            note.append('抄袭')

        # Check if not submitted
        if total_score == 0 and not note:
            note.append('未提交')

        ws_overall.append([
            student_id,
            eval.get('name', ''),
            total_score if student_id not in ORIGINAL_PLAGIARISM_STUDENTS else 0,  # 实验07-档位
            '',  # 实验-未完成 (预留列)
            total_score,  # 总分
            grade_label,  # 等级
            '; '.join(note)  # 备注
        ])

    for row in ws_overall.iter_rows(min_row=2):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="center")

    # Auto-adjust column widths for overall sheet
    overall_widths = [12, 12, 15, 15, 10, 10, 40]
    for i, width in enumerate(overall_widths, 1):
        ws_overall.column_dimensions[chr(64 + i)].width = width

    # ============================================================
    # Sheet 2: 单次实验评分（实验07-档位）
    # ============================================================
    ws1 = wb.create_sheet("单次实验评分")

    # Write headers (with quality assessment columns)
    headers = ['学号', '姓名', '团队协作(15)', '实验态度(10)', '完成度(35)',
               '代码质量(20)', '报告质量(20)', '质量评分', '总分', '等级', '抄袭警告', '备注(扣分)']
    ws1.append(headers)

    # Style header row
    for cell in ws1[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = border

    # Get quality scores and plagiarism data
    quality_scores = quality_data.get('quality_scores', {}) if quality_data else {}
    plagiarism_results = quality_data.get('plagiarism_data', {}).get('plagiarism_results', {}) if quality_data else {}
    suggestions = quality_data.get('suggestions', {}) if quality_data else {}

    # Use original plagiarism student list (from auto_score.py)
    # These students have been confirmed as plagiarists (>80% similarity)
    ORIGINAL_PLAGIARISM_STUDENTS = {
        '23071140217', '23071140216', '23071140214', '23071140228',
        '23071140233', '23071140220', '23071140223', '23071140213',
        '23071140219', '23071140204', '23071140208'
    }

    # Extract plagiarism students from suspicious pairs (for reference)
    suspicious_pairs = quality_data.get('plagiarism_data', {}).get('suspicious_pairs', []) if quality_data else []

    # Write student data
    for eval in evaluations:
        student_id = eval['student_id']
        quality = quality_scores.get(student_id, {})
        suggestion = suggestions.get(student_id, {})

        # Check for plagiarism
        plagiarism_warning = ""
        is_plagiarist = student_id in ORIGINAL_PLAGIARISM_STUDENTS

        if is_plagiarist:
            # Find similar students for this student
            similar_to = []
            for pair in suspicious_pairs:
                if pair.get('similarity', 0) > 80:
                    if pair.get('s1') == student_id:
                        similar_to.append((pair.get('s2', ''), pair.get('similarity', 0)))
                    elif pair.get('s2') == student_id:
                        similar_to.append((pair.get('s1', ''), pair.get('similarity', 0)))
            if similar_to:
                similar_str = ', '.join([f"{sid}({sim:.0f}%".replace('%', '') + '%)' for sid, sim in similar_to[:2]])
                plagiarism_warning = f"⚠️ 抄袭({similar_str})"
            else:
                plagiarism_warning = "⚠️ 抄袭"

        # For plagiarists, all scores are 0
        if is_plagiarist:
            team_score = 0
            attitude_score = 0
            completion_score = 0
            code_score = 0
            report_score = 0
            total = 0
            grade = '不及格'
        else:
            team_score = eval['scores']['team_collaboration']
            attitude_score = eval['scores']['attitude']
            completion_score = eval['scores']['completion']
            code_score = eval['scores']['code_quality']
            report_score = eval['scores']['report_quality']
            total = eval['total_score']
            grade = eval['grade_label']

        row = [
            student_id,
            eval.get('name', ''),
            team_score,
            attitude_score,
            completion_score,
            code_score,
            report_score,
            f"{quality.get('overall_quality', 0):.0f}" if not is_plagiarist else 0,
            total,
            grade,
            plagiarism_warning
        ]

        # Compile feedback notes (仅显示扣分/问题)
        deduction_notes = []

        # 检查各项得分是否满分，未满分则添加扣分说明
        max_scores = {
            'team_collaboration': 15,
            'attitude': 10,
            'completion': 35,
            'code_quality': 20,
            'report_quality': 20
        }

        for cat_id, max_score in max_scores.items():
            score = eval['scores'].get(cat_id, 0)
            if score < max_score:
                cat_names = {
                    'team_collaboration': '团队协作',
                    'attitude': '实验态度',
                    'completion': '完成度',
                    'code_quality': '代码质量',
                    'report_quality': '报告质量'
                }
                deduction = max_score - score
                deduction_notes.append(f"{cat_names[cat_id]}-{deduction}分")

        # 添加抄袭警告作为扣分
        if plagiarism_warning:
            deduction_notes.append(plagiarism_warning)

        row.append('; '.join(deduction_notes) if deduction_notes else '')

        ws1.append(row)

        # Style data cells
        for cell in ws1[ws1.max_row]:
            cell.border = border
            cell.alignment = Alignment(vertical="center")

        # Highlight plagiarism warnings
        if plagiarism_warning:
            ws1[f'K{ws1.max_row}'].fill = warning_fill

    # Auto-adjust column widths
    column_widths = [12, 12, 15, 15, 15, 15, 15, 10, 10, 10, 15, 50]
    for i, width in enumerate(column_widths, 1):
        ws1.column_dimensions[chr(64 + i)].width = width

    # 设置打印区域，将单次实验评分相关的所有表格放在同一页
    ws1.page_setup.printArea = f"$A$1:$L${ws1.max_row}"
    ws1.page_setup.fitToPage = True
    ws1.page_setup.fitToHeight = 0
    ws1.page_setup.fitToWidth = 1

    # ============================================================
    # Sheet 3: 详细评分
    # ============================================================
    ws2 = wb.create_sheet("详细评分")
    detail_headers = ['学号', '姓名', '评分项', '得分', '扣分说明']
    ws2.append(detail_headers)
    for cell in ws2[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment

    for eval in evaluations:
        student_id = eval['student_id']
        student_name = eval.get('name', '')
        for cat_id, score in eval['scores'].items():
            cat_info = next((c for c in rubric['categories'] if c['id'] == cat_id), None)
            cat_name = cat_info['name'] if cat_info else cat_id
            feedback = '; '.join(eval['feedback'].get(cat_id, []))
            ws2.append([student_id, student_name, cat_name, score, feedback])

    ws2.column_dimensions['A'].width = 12
    ws2.column_dimensions['B'].width = 12
    ws2.column_dimensions['C'].width = 15
    ws2.column_dimensions['D'].width = 10
    ws2.column_dimensions['E'].width = 60

    # ============================================================
    # Sheet 4: 统计分析
    # ============================================================
    ws3 = wb.create_sheet("统计分析")
    scores = [e['total_score'] for e in evaluations if e.get('total_score', 0) > 0]
    submitted_count = len(scores)
    missing_count = len(evaluations) - submitted_count

    ws3['A1'] = '统计项目'
    ws3['B1'] = '数值'
    ws3['A1'].font = header_font
    ws3['B1'].font = header_font

    stats = [
        ('班级总人数', len(evaluations)),
        ('提交报告', submitted_count),
        ('未提交', missing_count),
        ('平均分', round(sum(scores) / submitted_count, 2) if submitted_count > 0 else 0),
        ('最高分', max(scores) if submitted_count > 0 else 0),
        ('最低分', min(scores) if submitted_count > 0 else 0),
        ('及格率(>=60)', f"{sum(1 for s in scores if s >= 60) / submitted_count * 100:.1f}%" if submitted_count > 0 else "0%"),
        ('优秀率(>=90)', f"{sum(1 for s in scores if s >= 90) / submitted_count * 100:.1f}%" if submitted_count > 0 else "0%")
    ]

    for i, (label, value) in enumerate(stats, 2):
        ws3[f'A{i}'] = label
        ws3[f'B{i}'] = value

    ws3.column_dimensions['A'].width = 20
    ws3.column_dimensions['B'].width = 15

    # ============================================================
    # Sheet 5: 评估说明
    # ============================================================
    ws4 = wb.create_sheet("评估说明")

    title_font = Font(bold=True, size=14)
    ws4['A1'] = '自动评估说明'
    ws4['A1'].font = title_font

    notes = [
        "",
        "本评分表由系统自动生成，基于以下评估逻辑：",
        "",
        "1. 评估方法：关键词检测",
        "   - 系统检测报告中是否包含各个章节和关键技术关键词",
        "   - 存在关键词即给予相应分数，未完全评估内容质量",
        "",
        "2. 建议调整事项：",
        "   - 【团队协作分】：建议根据实际分工情况调整",
        "   - 【实验态度分】：需根据考勤记录核实后修正",
        "   - 【实验完成度】：建议对照实际演示/验收结果调整",
        "   - 【代码质量】：建议人工审阅代码后调整",
        "   - 【报告质量】：建议人工审阅报告质量后调整",
        "",
        "3. 使用建议：",
        "   - 自动评分仅作为参考基准",
        "   - 教师应人工审阅后确定最终分数",
        "   - 可直接在Excel中修改分数",
        "",
        "4. 未提交学生：",
        "   - 5名学生未提交报告，记0分",
        "   - 学号: 23071140205, 23071140209, 23071140229, 23071140234, 23071140235"
    ]

    for i, note in enumerate(notes, 1):
        ws4[f'A{i}'] = note

    ws4.column_dimensions['A'].width = 80

    # Save
    output_path = TEACHER_OUTPUT / "汽服2302B班_07_档位实验_评分表.xlsx"
    wb.save(output_path)
    print(f"  Saved: {output_path}")

def create_student_feedback(evaluations, rubric, quality_data=None):
    """Create feedback documents for students"""
    print("Creating student feedback documents...")

    # Original plagiarism student list
    ORIGINAL_PLAGIARISM_STUDENTS = {
        '23071140217', '23071140216', '23071140214', '23071140228',
        '23071140233', '23071140220', '23071140223', '23071140213',
        '23071140219', '23071140204', '23071140208'
    }

    # Get suspicious pairs for plagiarism details
    suspicious_pairs = quality_data.get('plagiarism_data', {}).get('suspicious_pairs', []) if quality_data else []

    # Get reference answers and common issues
    ref_answers = rubric.get('reference_answers', {})
    common_issues = rubric.get('common_issues', {})

    # Get personalized suggestions
    suggestions = quality_data.get('suggestions', {}) if quality_data else {}
    quality_scores = quality_data.get('quality_scores', {}) if quality_data else {}

    # Create feedback template (only for students who submitted)
    for eval in evaluations:
        # Skip missing submissions
        if eval.get('total_score', 0) == 0 and eval.get('feedback', {}).get('overall') == ['未提交实验报告']:
            print(f"  Skipping {eval['student_id']}: No report submitted")
            continue

        student_id = eval['student_id']
        student_suggestions = suggestions.get(student_id, {})
        quality = quality_scores.get(student_id, {})

        doc = Document()

        # Title
        title = doc.add_paragraph('实验报告反馈')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(18)
        title.runs[0].font.bold = True

        # Experiment info
        doc.add_paragraph(f"实验名称: {rubric['experiment_name']}")
        doc.add_paragraph(f"学号: {student_id}")

        # Check for plagiarism
        is_plagiarist = student_id in ORIGINAL_PLAGIARISM_STUDENTS

        # Find similar students for this student
        similar_students = []
        for pair in suspicious_pairs:
            if pair.get('similarity', 0) > 60:  # Show all similarities >60%
                if pair.get('s1') == student_id:
                    similar_students.append((pair.get('s2', ''), pair.get('similarity', 0)))
                elif pair.get('s2') == student_id:
                    similar_students.append((pair.get('s1', ''), pair.get('similarity', 0)))

        # Sort by similarity descending
        similar_students.sort(key=lambda x: -x[1])

        # Plagiarism warning if plagiarist or high similarity
        if is_plagiarist:
            warning_para = doc.add_paragraph()
            warning_run = warning_para.add_run("⚠️ 抄袭判定: ")
            warning_run.font.color.rgb = RGBColor(255, 0, 0)
            warning_run.font.bold = True
            warning_para.add_run("你的报告被判定为抄袭，本次实验成绩为0分。")
            doc.add_paragraph()
        elif similar_students and similar_students[0][1] > 70:
            # Show warning for high similarity but not plagiarist
            warning_para = doc.add_paragraph()
            warning_run = warning_para.add_run("⚠️ 相似度警告: ")
            warning_run.font.color.rgb = RGBColor(255, 165, 0)
            warning_run.font.bold = True
            warning_para.add_run("你的报告内容与以下同学高度相似，请注意原创性。")
            doc.add_paragraph()

        # Show similar students if any
        if similar_students:
            doc.add_heading("相似度检测", level=2)
            for sid, sim in similar_students[:5]:  # Show top 5
                para = doc.add_paragraph()
                para.add_run(f"• 学号 {sid}: ").bold = True
                if sim >= 90:
                    sim_status = "极高"
                    color = RGBColor(255, 0, 0)
                elif sim >= 80:
                    sim_status = "很高"
                    color = RGBColor(255, 0, 0)
                elif sim >= 70:
                    sim_status = "较高"
                    color = RGBColor(255, 165, 0)
                else:
                    sim_status = "中等"
                    color = RGBColor(0, 0, 255)
                run = para.add_run(f"{sim:.1f}% ({sim_status}相似)")
                run.font.color.rgb = color
            doc.add_paragraph()

        # Grade
        grade_para = doc.add_paragraph()
        grade_para.add_run("评价等级: ").font.size = Pt(14)

        if is_plagiarist:
            grade_text = "抄袭 (0分)"
            grade_color = RGBColor(255, 0, 0)
        else:
            grade_text = f"{eval['grade_label']}"
            grade_color = RGBColor(0, 128, 0) if eval['grade'] in ['A', 'B', 'C'] else RGBColor(255, 0, 0)

        grade_run = grade_para.add_run(grade_text)
        grade_run.font.size = Pt(16)
        grade_run.font.bold = True
        grade_run.font.color.rgb = grade_color

        doc.add_paragraph()  # Blank line

        # Detailed score breakdown (skip for plagiarists)
        if is_plagiarist:
            doc.add_heading("抄袭详情", level=2)
            doc.add_paragraph("由于报告被判定为抄袭，所有评分项目均为0分。")
            doc.add_paragraph("如对判定有异议，请联系教师申诉。")

            # Still show similar students for plagiarists
            if similar_students:
                doc.add_paragraph()
                doc.add_heading("相似同学列表", level=2)
                for sid, sim in similar_students[:5]:
                    para = doc.add_paragraph()
                    para.add_run(f"• 学号 {sid}: ").bold = True
                    run = para.add_run(f"{sim:.1f}%")
                    run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            # Show detailed scores for non-plagiarists
            doc.add_heading("详细得分", level=2)
            scores = eval['scores']
            total = eval['total_score']

            # Create score table
            score_items = [
                ("团队协作", scores['team_collaboration'], 5),
                ("实验态度", scores['attitude'], 10),
                ("实验原理", scores['principle_understanding'], 10),
                ("实验完成度", scores['completion'], 35),
                ("代码质量", scores['code_quality'], 30),
                ("报告质量", scores['report_quality'], 10),
            ]

            for item_name, score, max_score in score_items:
                para = doc.add_paragraph()
                para.add_run(f"{item_name}: ").bold = True
            percentage = (score / max_score * 100) if max_score > 0 else 0
            if percentage >= 90:
                status = "优秀"
                color = RGBColor(0, 128, 0)
            elif percentage >= 70:
                status = "良好"
                color = RGBColor(0, 0, 255)
            elif percentage >= 60:
                status = "及格"
                color = RGBColor(255, 165, 0)
            else:
                status = "需改进"
                color = RGBColor(255, 0, 0)
            run = para.add_run(f"{score}/{max_score}分 ({status})")
            run.font.color.rgb = color

        doc.add_paragraph()
        para = doc.add_paragraph()
        para.add_run(f"总分: {total}/100分  等级: {eval['grade_label']}").bold = True

        doc.add_paragraph()  # Blank line

        # Detailed feedback for each category
        doc.add_heading("各部分评价", level=2)

        feedback = eval.get('feedback', {})

        # Team collaboration
        doc.add_heading("1. 团队协作", level=3)
        team_feedback = feedback.get('team_collaboration', [])
        if team_feedback:
            for item in team_feedback:
                doc.add_paragraph(f"• {item}")
        else:
            doc.add_paragraph("团队协作信息完整，分工明确。")

        # Principle understanding
        doc.add_heading("2. 实验原理与认知", level=3)
        principle_feedback = feedback.get('principle_understanding', [])
        if principle_feedback:
            for item in principle_feedback:
                doc.add_paragraph(f"• {item}")
        else:
            doc.add_paragraph("实验原理阐述准确，目的清晰。")

        # Completion
        doc.add_heading("3. 实验完成度", level=3)
        completion_feedback = feedback.get('completion', [])
        if completion_feedback:
            for item in completion_feedback:
                doc.add_paragraph(f"• {item}")
        else:
            doc.add_paragraph("实验完成度良好，结果记录完整。")

        # Code quality
        doc.add_heading("4. 代码质量", level=3)
        code_feedback = feedback.get('code_quality', [])
        if code_feedback:
            for item in code_feedback:
                doc.add_paragraph(f"• {item}")
        else:
            doc.add_paragraph("代码结构清晰，注释详尽。")

        # Report quality
        doc.add_heading("5. 报告质量", level=3)
        report_feedback = feedback.get('report_quality', [])
        if report_feedback:
            for item in report_feedback:
                doc.add_paragraph(f"• {item}")
        else:
            doc.add_paragraph("报告格式规范，内容完整。")

        doc.add_paragraph()  # Blank line

        # Technical guidance based on weak areas
        weak_areas = []
        if scores['completion'] < 28:
            weak_areas.append(("实验完成度", "检查硬件接线图、GPIO配置说明、实验现象记录是否完整"))

        if scores['code_quality'] < 24:
            weak_areas.append(("代码质量", "增加代码注释，说明DWT消抖、中断回调、状态机实现逻辑"))

        if scores['report_quality'] < 8:
            weak_areas.append(("报告质量", "补充问题讨论、个人心得和思考题回答"))

        if scores['principle_understanding'] < 8:
            weak_areas.append(("实验原理", "详细说明外部中断、消抖方法、汽车电子应用场景"))

        if weak_areas:
            doc.add_heading("改进建议", level=2)
            for area_name, suggestion in weak_areas:
                doc.add_paragraph()
                doc.add_paragraph(f"🔧 {area_name}:", style='List Bullet').runs[0].bold = True
                doc.add_paragraph(f"   {suggestion}")

        doc.add_paragraph()  # Blank line

        # Positive feedback for good performance
        good_areas = []
        if scores['completion'] >= 32:
            good_areas.append("实验完成度很高，现象记录详细")
        if scores['code_quality'] >= 27:
            good_areas.append("代码质量优秀，注释规范完整")
        if scores['report_quality'] >= 9:
            good_areas.append("报告质量高，内容充实完整")
        if scores['principle_understanding'] >= 9:
            good_areas.append("对实验原理理解深入")

        if good_areas:
            doc.add_heading("报告亮点", level=2)
            for item in good_areas:
                doc.add_paragraph(f"✓ {item}")

        # Save document
        output_path = STUDENT_OUTPUT / f"{eval['student_id']}_反馈.docx"
        doc.save(output_path)

    print(f"  Created {len(evaluations)} feedback documents")

def main():
    print("Generating output files...")

    # Create output directories
    TEACHER_OUTPUT.mkdir(parents=True, exist_ok=True)
    STUDENT_OUTPUT.mkdir(parents=True, exist_ok=True)

    # Load evaluations
    eval_path = PROCESSED_DIR / "evaluations.json"
    if not eval_path.exists():
        print("Error: evaluations.json not found. Run evaluate.py first.")
        return

    with open(eval_path, 'r', encoding='utf-8') as f:
        evaluations = json.load(f)

    # Load rubric
    rubric_path = DATA_DIR / "rubric.json"
    with open(rubric_path, 'r', encoding='utf-8') as f:
        rubric = json.load(f)

    # Load quality assessment data if available
    quality_path = PROCESSED_DIR / "quality_assessment.json"
    quality_data = None
    if quality_path.exists():
        with open(quality_path, 'r', encoding='utf-8') as f:
            quality_data = json.load(f)
        print("Using quality assessment data...")

    # Generate outputs
    create_teacher_excel(evaluations, rubric, quality_data)
    create_student_feedback(evaluations, rubric, quality_data)

    print("\nOutput generation complete!")
    print(f"Teacher Excel: {TEACHER_OUTPUT}")
    print(f"Student Feedbacks: {STUDENT_OUTPUT}")

    # Print plagiarism summary if available
    if quality_data and quality_data.get('plagiarism_data', {}).get('suspicious_pairs'):
        suspicious = quality_data['plagiarism_data']['suspicious_pairs']
        print(f"\nWARNING: Found {len(suspicious)} suspicious report pairs (>60% similarity)")
        print("  These have been marked in the Excel file.")

if __name__ == "__main__":
    main()
