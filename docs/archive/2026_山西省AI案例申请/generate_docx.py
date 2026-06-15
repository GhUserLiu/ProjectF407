#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""使用python-docx生成Word文档（可另存为PDF）"""

import sys
import re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 读取Markdown文件
with open('创AI案例-开发与应用报告.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# 创建Word文档
doc = Document()

# 设置中文字体
def set_chinese_font(run):
    """设置中文字体"""
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 当前函数
def add_title(text, level=1):
    """添加标题"""
    sizes = {1: 18, 2: 16, 3: 14, 4: 12}
    size = sizes.get(level, 11)

    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.size = Pt(size)
    run.font.bold = True
    set_chinese_font(run)

def add_paragraph(text):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    set_chinese_font(run)

def add_code(code_text):
    """添加代码块"""
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    # 设置灰色背景
    shading = run._element.get_or_add_oxml_element('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')

def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid'

    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        # 设置表头样式
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                set_chinese_font(run)

    # 表格内容
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            cell = row_cells[j]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    set_chinese_font(run)

def add_list(items, ordered=False):
    """添加列表"""
    for i, item in enumerate(items):
        para = doc.add_paragraph()
        if ordered:
            run = para.add_run(f'{i+1}. ')
        else:
            run = para.add_run('• ')
        run.font.bold = True
        set_chinese_font(run)

        run = para.add_run(item)
        run.font.size = Pt(11)
        set_chinese_font(run)

# 添加封面
title = doc.add_heading('基于多算法融合的STM32实验报告智能评估系统', 0)
run = title.runs[0]
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1a, 0x54, 0x90)
set_chinese_font(run)

doc.add_paragraph('开发与应用报告', style='Heading 2')

# 元数据
metadata = doc.add_paragraph()
metadata.add_run('申报类别：创AI - 智能信息系统\n')
metadata.add_run('申报教师：刘兆骐\n')
metadata.add_run('学校：山西工程科技职业大学\n')
metadata.add_run('联系方式：19335411556\n')
metadata.add_run('申报日期：2026年6月12日')
for run in metadata.runs:
    set_chinese_font(run)

doc.add_page_break()

# 解析Markdown内容
lines = md_content.split('\n')
in_code_block = False
code_buffer = []
list_items = []
is_ordered_list = False

i = 0
while i < len(lines):
    line = lines[i]

    # 跳过封面部分（已手动添加）
    if i < 20:
        if line.startswith('# '):
            break
        i += 1
        continue

    # 代码块处理
    if line.startswith('```'):
        if not in_code_block:
            in_code_block = True
        else:
            add_code('\n'.join(code_buffer))
            code_buffer = []
            in_code_block = False
        i += 1
        continue

    if in_code_block:
        code_buffer.append(line)
        i += 1
        continue

    # 标题处理
    if line.startswith('#'):
        match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if match:
            level = len(match.group(1))
            text = match.group(2)
            add_title(text, level)
            i += 1
            continue

    # 表格处理
    if line.startswith('|'):
        if i + 1 < len(lines) and lines[i + 1].startswith('|---'):
            headers = [h.strip() for h in line.split('|')[1:-1]]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                cells = [c.strip() for c in lines[i].split('|')[1:-1]]
                rows.append(cells)
                i += 1
            add_table(headers, rows)
            continue

    # 列表处理
    if line.strip().startswith('- '):
        list_items = []
        is_ordered_list = False
        while i < len(lines) and lines[i].strip().startswith('- '):
            list_items.append(lines[i].strip()[2:])
            i += 1
        add_list(list_items, ordered=False)
        continue

    if re.match(r'^\d+\. ', line.strip()):
        list_items = []
        is_ordered_list = True
        while i < len(lines) and re.match(r'^\d+\. ', lines[i].strip()):
            content = re.sub(r'^\d+\. ', '', lines[i].strip())
            list_items.append(content)
            i += 1
        add_list(list_items, ordered=True)
        continue

    # 分隔线
    if line.strip() in ['---', '***']:
        doc.add_paragraph('_' * 80)
        i += 1
        continue

    # 空行
    if line.strip() == '':
        doc.add_paragraph()
        i += 1
        continue

    # 普通段落
    if line.strip():
        # 处理粗体
        para = doc.add_paragraph()

        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))', line.strip())
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            elif re.match(r'\[([^\]]+)\]\([^\)]+\)', part):
                # 链接
                text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', part)
                run = para.add_run(text)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                run.font.underline = True
            else:
                run = para.add_run(part)
            set_chinese_font(run)

    i += 1

# 保存Word文档
output_file = '创AI案例-开发与应用报告.docx'
doc.save(output_file)

print(f'Word文档生成成功: {output_file}')
print('请在Word中打开该文档，然后选择"文件" > "另存为" > 选择PDF格式')
