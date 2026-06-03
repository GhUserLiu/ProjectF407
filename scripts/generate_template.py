# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='宋体', size=12, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_paragraph_spacing(paragraph, line_spacing=1.35, space_before=0, space_after=0):
    """设置段落间距"""
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Cm(space_before)
    paragraph.paragraph_format.space_after = Cm(space_after)
    # 取消自动编号
    paragraph.paragraph_format.left_indent = Cm(0)

# 创建新文档
doc = Document()

# 设置页面边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# 设置默认正文样式（宋体、12pt、1.35倍行距）
normal_style = doc.styles['Normal']
normal_style.font.name = '宋体'
normal_style.font.size = Pt(12)
normal_style.paragraph_format.line_spacing = 1.35

# 设置一级标题样式
h1 = doc.styles['Heading 1']
h1.font.name = '黑体'
h1.font.size = Pt(16)
h1.paragraph_format.line_spacing = 1.35
h1.font.bold = True

# 设置二级标题样式
h2 = doc.styles['Heading 2']
h2.font.name = '黑体'
h2.font.size = Pt(14)
h2.paragraph_format.line_spacing = 1.35
h2.font.bold = True

# 标题页
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('汽车微处理器原理与应用')
set_chinese_font(run, '黑体', 22, True)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.0, 0, 0)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('实　验　报　告')
set_chinese_font(run, '黑体', 18, True)

# 添加多个空行
for _ in range(4):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 1.0, 0, 0)

# 基本信息
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('基本信息')
set_chinese_font(run, '黑体', 16, True)
set_paragraph_spacing(p, 1.35, 0.5, 0.3)

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'

info = [
    ['实验名称', '', '学生姓名', ''],
    ['学　　号', '', '班　　级', ''],
    ['组　　号', '', '指导教师', ''],
    ['实验日期', '', '提交日期', '']
]

for i, row_data in enumerate(info):
    for j, text in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_chinese_font(run, '宋体', 12)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

# 一、团队信息与分工
p = doc.add_paragraph()
run = p.add_run('一、团队信息与分工（5分）')
set_chinese_font(run, '黑体', 16, True)
set_paragraph_spacing(p, 1.35, 0.5, 0.3)

p = doc.add_paragraph()
run = p.add_run('1.1 团队成员基本信息')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
headers = ['组号', '角色', '姓名', '学号']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            set_chinese_font(run, '宋体', 12)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.3)

p = doc.add_paragraph()
run = p.add_run('1.2 个人分工说明')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请描述本人在团队中承担的具体工作任务')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.3)

p = doc.add_paragraph()
run = p.add_run('1.3 团队协作过程记录')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请记录团队讨论、问题解决、任务分配等协作过程')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

# 二、实验目的与原理
p = doc.add_paragraph()
run = p.add_run('二、实验目的与原理（10分）')
set_chinese_font(run, '黑体', 16, True)
set_paragraph_spacing(p, 1.35, 0.5, 0.3)

p = doc.add_paragraph()
run = p.add_run('2.1 实验目的')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请列出本次实验所要达到的目的和预期学习成果')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0.3)

p = doc.add_paragraph()
run = p.add_run('1. ')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

p = doc.add_paragraph()
run = p.add_run('2. ')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

p = doc.add_paragraph()
run = p.add_run('3. ')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.3)

p = doc.add_paragraph()
run = p.add_run('2.2 实验原理')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：正确阐述实验所涉及的基本原理、技术背景和相关理论知识')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.3)

p = doc.add_paragraph()
run = p.add_run('2.3 汽车电子应用场景')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：说明本实验在实际汽车电子系统中的应用背景和意义')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

# 三、硬件设计与连接
p = doc.add_paragraph()
run = p.add_run('三、硬件设计与连接（15分）')
set_chinese_font(run, '黑体', 16, True)
set_paragraph_spacing(p, 1.35, 0.5, 0.3)

p = doc.add_paragraph()
run = p.add_run('3.1 硬件连接图')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请绘制或插入硬件连接图，清晰标注各引脚连接关系')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0.5)

p = doc.add_paragraph()
run = p.add_run('3.2 引脚配置说明')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
headers = ['功能', 'GPIO引脚', '配置说明']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            set_chinese_font(run, '宋体', 12)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.3)

p = doc.add_paragraph()
run = p.add_run('说明：请阐述各引脚配置的理由和功能')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

p = doc.add_paragraph()
run = p.add_run('3.3 ISP烧录电路说明')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请说明ISP烧录电路的连接方式和烧录工具使用方法')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

# 四、软件设计与实现
p = doc.add_paragraph()
run = p.add_run('四、软件设计与实现（30分）')
set_chinese_font(run, '黑体', 16, True)
set_paragraph_spacing(p, 1.35, 0.5, 0.3)

p = doc.add_paragraph()
run = p.add_run('4.1 核心代码流程图')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请绘制核心程序的流程图，清晰展示程序逻辑结构')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0.5)

p = doc.add_paragraph()
run = p.add_run('4.2 关键代码片段及注释')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请粘贴关键代码片段，并添加详细注释说明')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0.5)

p = doc.add_paragraph()
run = p.add_run('4.3 中断服务程序说明')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请说明中断服务程序的配置方法、触发条件和处理逻辑')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

p = doc.add_paragraph()
run = p.add_run('4.4 团队成员负责的代码模块说明')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
headers = ['成员姓名', '负责模块', '完成情况']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            set_chinese_font(run, '宋体', 12)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

# 五、实验结果分析
p = doc.add_paragraph()
run = p.add_run('五、实验结果分析（20分）')
set_chinese_font(run, '黑体', 16, True)
set_paragraph_spacing(p, 1.35, 0.5, 0.3)

p = doc.add_paragraph()
run = p.add_run('5.1 实验现象记录')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请详细记录实验过程中的各种现象')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.3)

p = doc.add_paragraph()
run = p.add_run('5.2 结果照片或视频截图')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请插入实验结果照片或视频截图，并适当标注说明')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0.5)

p = doc.add_paragraph()
run = p.add_run('5.3 与预期结果对比分析')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请对比实验结果与预期结果，分析差异原因')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

# 六、问题与讨论
p = doc.add_paragraph()
run = p.add_run('六、问题与讨论（15分）')
set_chinese_font(run, '黑体', 16, True)
set_paragraph_spacing(p, 1.35, 0.5, 0.3)

p = doc.add_paragraph()
run = p.add_run('6.1 调试过程中遇到的问题')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)

table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'
headers = ['问题现象', '可能原因', '解决方法']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            set_chinese_font(run, '宋体', 12)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.3)

p = doc.add_paragraph()
run = p.add_run('6.2 团队协作解决过程')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请描述团队如何协作解决上述问题')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.3)

p = doc.add_paragraph()
run = p.add_run('6.3 个人心得体会（本部分需独立撰写）')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)
p = doc.add_paragraph()
run = p.add_run('说明：请总结本次实验的个人收获、遇到的问题及解决过程，谈谈个人体会和感悟')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

# 七、思考题回答
p = doc.add_paragraph()
run = p.add_run('七、思考题回答（5分）')
set_chinese_font(run, '黑体', 16, True)
set_paragraph_spacing(p, 1.35, 0.5, 0.3)
p = doc.add_paragraph()
run = p.add_run('说明：请回答实验任务书中的思考题')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0.3)

for i in range(1, 8):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    set_chinese_font(p.runs[0])
    set_paragraph_spacing(p, 1.35, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run('   答：')
    set_chinese_font(p.runs[0])
    set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

# 八、附录：核心代码
p = doc.add_paragraph()
run = p.add_run('八、附录：核心代码')
set_chinese_font(run, '黑体', 16, True)
set_paragraph_spacing(p, 1.35, 0.5, 0.3)
p = doc.add_paragraph()
run = p.add_run('说明：请在此处粘贴完整的核心代码，代码需有详细注释')
set_chinese_font(p.runs[0])
set_paragraph_spacing(p, 1.35, 0, 0.5)

# 添加代码示例框
p = doc.add_paragraph()
p.add_run('在此处粘贴核心代码...')
set_chinese_font(p.runs[0], 'Consolas', 10.5)
set_paragraph_spacing(p, 1.0, 0, 0.5)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

# 报告要求
p = doc.add_paragraph()
run = p.add_run('报告要求')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)

requirements = [
    '1. 报告格式：Word（.docx）或 WPS 文字格式',
    '2. 正文字体：宋体，小四号（12pt），1.35倍行距',
    '3. 代码格式：等宽字体（Consolas/Courier New），关键部分需加注释',
    '4. 图表要求：清晰规范，有图题和编号',
    '5. 照片要求：嵌入文档，适当标注说明',
    '6. 字数要求：报告总字数不少于2000字',
    '7. 核心代码：必须附在报告附录中，需有详细注释',
    '8. 心得体会：必须独立撰写，体现个人思考'
]

for req in requirements:
    p = doc.add_paragraph()
    run = p.add_run(req)
    set_chinese_font(run, '宋体', 12)
    set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

# 提交说明
p = doc.add_paragraph()
run = p.add_run('提交说明')
set_chinese_font(run, '黑体', 14, True)
set_paragraph_spacing(p, 1.35, 0.3, 0.2)

submissions = [
    '提交平台：学习通（超星学习平台）',
    '文件命名：组号-姓名.docx（例如：第1组-张三.docx）',
    '提交方式：每位成员单独提交',
    '内容说明：同组成员可提交完全一致的报告内容，仅"个人心得体会"部分需独立撰写',
    '提交时间：按教师要求准时提交',
    '重要提醒：逾期提交将按迟交处理（每天扣5分）'
]

for sub in submissions:
    p = doc.add_paragraph()
    run = p.add_run(sub)
    set_chinese_font(run, '宋体', 12)
    set_paragraph_spacing(p, 1.35, 0, 0)

doc.add_paragraph()
set_paragraph_spacing(doc.paragraphs[-1], 1.35, 0, 0.5)

# 页脚信息
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('模板编制：汽车微处理器原理与应用课程组 | 编制日期：2026年6月')
set_chinese_font(run, '宋体', 10)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
set_paragraph_spacing(p, 1.35, 0.5, 0)

# 保存
doc.save('c:/Users/liuzh/Projects/Workspace/stm32f407/docs/实验报告模板.docx')
print('实验报告模板已重新生成')
print('修复：小标题编号不再使用自动编号，改为手动输入')
